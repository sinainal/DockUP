from __future__ import annotations

import html
import inspect
import json
import os
import re
import signal
import shutil
import subprocess
import sys
import time
from pathlib import Path
from typing import Any

from fastapi import APIRouter, BackgroundTasks, HTTPException
from fastapi.responses import FileResponse, JSONResponse

from ..config import BASE, DATA_DIR, DOCK_DIR, WORKSPACE_DIR
from ..helpers import (
    normalize_docking_config,
    read_json,
    relative_to_base,
    resolve_dock_directory,
    to_display_path,
    timestamp_token,
    write_json,
    BASE_RESOLVED,
    DOCK_DIR_RESOLVED,
    WORKSPACE_RESOLVED,
)
from ..models import GraphPayload, RenderPayload, ReportCompilePayload
from ..state import REPORT_STATE

router = APIRouter()

REPORT_PREDEFINED_PLOTS: dict[str, dict[str, str]] = {
    "affinity_table_plus_boxplot": {
        "label": "Affinity Table + Boxplot",
        "module": "figure_scripts.final_plots.affinity_variants",
        "filename": "affinity_boxplot.png",
    },
    "interaction_frequency_heatmap": {
        "label": "Interaction Frequency Heatmap",
        "module": "figure_scripts.final_plots.interacted_residue_plots",
        "filename": "run_frequency_heatmap.png",
    },
    "common_residue_heatmap": {
        "label": "Common Residue Heatmap",
        "module": "figure_scripts.final_plots.common_residue_interactions",
        "filename": "common_residue_heatmap.png",
    },
    "interaction_stacked_bar": {
        "label": "Interaction Stacked Bar",
        "module": "figure_scripts.final_plots.interaction_plots",
        "filename": "interaction_stacked_bar.png",
    },
}

REPORT_PANEL_DTYPES: tuple[str, ...] = ("D1", "D2", "D3", "D4", "D5")
REPORT_DTYPE_RE = re.compile(r"^D\d+$", re.IGNORECASE)
REPORT_RUN_RE = re.compile(r"^run(\d+)$", re.IGNORECASE)
REPORT_CASE_RUN_RE = re.compile(r"^([A-Za-z0-9]+)_(.+)_run(\d+)$", re.IGNORECASE)
REPORT_MAX_DISCOVERY_DEPTH = 6
REPORT_IMAGE_EXTENSIONS: tuple[str, ...] = (".png", ".jpg", ".jpeg", ".webp", ".svg")
REPORT_METADATA_FILENAME = ".docking_app_meta.json"
REPORT_IMAGE_METADATA_SUFFIX = ".meta.json"
REPORT_TEMPLATE_HEADINGS: tuple[str, ...] = (
    "Materials and Methods: Molecular Docking Simulations",
    "Results",
    "Discussion",
    "Conclusion",
)
REPORT_PLOT_ORDER_BY_NAME: tuple[str, ...] = (
    "affinity_boxplot",
    "run_frequency_heatmap",
    "common_residue_heatmap",
    "interaction_stacked_bar",
)
REPORT_RENDER_MODE_CLASSIC = "classic"
REPORT_RENDER_MODE_OTOFIGURE = "otofigure"
REPORT_RENDER_MODE_MULTI_LIGAND = "multi_ligand_panel"


def _prettify_label(name: str, *, trim_run_suffix: bool = False) -> str:
    text = str(name or "").strip()
    if trim_run_suffix:
        text = re.sub(r"_\d+$", "", text)
    text = re.sub(r"[_-]+", " ", text).strip()
    return text or str(name or "")


def _normalize_render_mode(raw_value: Any) -> str:
    value = str(raw_value or "").strip().lower()
    if not value or value in {"classic", "panel", "grid", "default"}:
        return REPORT_RENDER_MODE_CLASSIC
    if value in {"otofigure", "multiview", "multi_view", "multi-run", "multi_run"}:
        return REPORT_RENDER_MODE_OTOFIGURE
    if value in {"multi_ligand", "multi-ligand", "multi_ligand_panel", "multi-ligand-panel"}:
        return REPORT_RENDER_MODE_MULTI_LIGAND
    raise ValueError(f"Unsupported render mode: {raw_value}")


def _receptor_sort_key(name: str) -> tuple[int, int, str]:
    match = REPORT_DTYPE_RE.fullmatch(name or "")
    if match:
        try:
            return (0, int((name or "")[1:]), (name or "").upper())
        except ValueError:
            return (0, 10**9, (name or "").upper())
    return (1, 10**9, (name or "").lower())


def _ligand_sort_key(name: str) -> tuple[str, str]:
    pretty = _prettify_label(name, trim_run_suffix=True).lower()
    return (pretty, (name or "").lower())


def _metadata_file_for_source(source_dir: Path) -> Path:
    return source_dir.resolve() / REPORT_METADATA_FILENAME


def _image_metadata_path(image_path: Path) -> Path:
    return image_path.with_name(f"{image_path.name}{REPORT_IMAGE_METADATA_SUFFIX}")


def _read_image_metadata(image_path: Path) -> dict[str, Any]:
    raw = read_json(_image_metadata_path(image_path), {})
    return raw if isinstance(raw, dict) else {}


def _write_image_metadata(image_path: Path, payload: dict[str, Any]) -> None:
    write_json(_image_metadata_path(image_path), payload)


def _delete_image_artifacts(image_path: Path) -> None:
    image_path.unlink(missing_ok=True)
    _image_metadata_path(image_path).unlink(missing_ok=True)


def _collect_source_entities(source_dir: Path) -> tuple[list[str], list[str]]:
    inventory = _collect_receptor_inventory(source_dir)
    candidates = _collect_receptor_candidates(source_dir)
    receptor_ids = sorted(set(inventory.keys()) | set(candidates.keys()), key=_receptor_sort_key)

    ligand_names: set[str] = set()
    for ligand_map in inventory.values():
        ligand_names.update(ligand_map.keys())
    for ligand_map in candidates.values():
        ligand_names.update(ligand_map.keys())

    ligands = sorted(ligand_names, key=_ligand_sort_key)
    return receptor_ids, ligands


def _default_main_type_label(source_dir: Path) -> str:
    default = _prettify_label(source_dir.name, trim_run_suffix=False)
    lowered = default.lower()
    if lowered in {
        "data",
        "dock",
        "results",
        "dimer final",
        "dimer final linked",
        "dimer full",
        "report outputs",
    }:
        return ""
    return default


def _normalize_label_map(raw_map: Any) -> dict[str, str]:
    out: dict[str, str] = {}
    if not isinstance(raw_map, dict):
        return out
    for key, value in raw_map.items():
        raw_key = str(key or "").strip()
        raw_value = str(value or "").strip()
        if not raw_key:
            continue
        if raw_value:
            out[raw_key] = raw_value
    return out


def _normalize_positive_int(raw_value: Any, *, default: int = 1, min_value: int = 1, max_value: int = 10000) -> int:
    try:
        val = int(raw_value)
    except (TypeError, ValueError):
        val = default
    val = max(min_value, min(max_value, val))
    return val


def _normalize_extra_sections(raw_sections: Any) -> list[dict[str, str]]:
    out: list[dict[str, str]] = []
    if not isinstance(raw_sections, list):
        return out
    for item in raw_sections:
        if not isinstance(item, dict):
            continue
        title = str(item.get("title") or "").strip()
        body = str(item.get("body") or "").strip()
        if not title and not body:
            continue
        out.append({"title": title, "body": body})
    return out


def _normalize_caption_map(raw_map: Any) -> dict[str, str]:
    out: dict[str, str] = {}
    if not isinstance(raw_map, dict):
        return out
    for key, value in raw_map.items():
        k = str(key or "").strip()
        v = str(value or "").strip()
        if not k or not v:
            continue
        out[k] = v
    return out


def _normalize_order_list(raw_list: Any, allowed_items: list[str]) -> list[str]:
    allowed = [str(item) for item in allowed_items if str(item)]
    allowed_set = set(allowed)
    out: list[str] = []
    seen: set[str] = set()

    if isinstance(raw_list, list):
        for item in raw_list:
            key = str(item or "").strip()
            if not key or key in seen or key not in allowed_set:
                continue
            seen.add(key)
            out.append(key)

    for key in allowed:
        if key in seen:
            continue
        seen.add(key)
        out.append(key)
    return out


def _load_source_metadata(
    source_dir: Path,
    receptor_ids: list[str] | None = None,
    ligand_names: list[str] | None = None,
) -> dict[str, Any]:
    source_dir = source_dir.resolve()
    if receptor_ids is None or ligand_names is None:
        detected_receptors, detected_ligands = _collect_source_entities(source_dir)
        if receptor_ids is None:
            receptor_ids = detected_receptors
        if ligand_names is None:
            ligand_names = detected_ligands
    receptor_ids = list(receptor_ids)
    ligand_names = list(ligand_names)
    meta_file = _metadata_file_for_source(source_dir)

    raw: dict[str, Any] = {}
    if meta_file.exists() and meta_file.is_file():
        try:
            raw = json.loads(meta_file.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            raw = {}

    receptor_order = _normalize_order_list(raw.get("receptor_order"), receptor_ids)
    ligand_order = _normalize_order_list(raw.get("ligand_order"), ligand_names)

    receptor_override = _normalize_label_map(raw.get("receptor_labels"))
    ligand_override = _normalize_label_map(raw.get("ligand_labels"))

    receptor_labels: dict[str, str] = {}
    for receptor_id in receptor_order:
        fallback = _prettify_label(receptor_id, trim_run_suffix=False)
        receptor_labels[receptor_id] = receptor_override.get(receptor_id, fallback)

    ligand_labels: dict[str, str] = {}
    for ligand_name in ligand_order:
        fallback = _prettify_label(ligand_name, trim_run_suffix=True)
        ligand_labels[ligand_name] = ligand_override.get(ligand_name, fallback)

    report_raw = raw.get("report") if isinstance(raw.get("report"), dict) else {}
    figure_start_number = _normalize_positive_int(
        report_raw.get("figure_start_number", raw.get("figure_start_number", 1)),
        default=1,
        min_value=1,
        max_value=999,
    )
    extra_sections = _normalize_extra_sections(report_raw.get("extra_sections", raw.get("extra_sections", [])))
    figure_caption_overrides = _normalize_caption_map(
        report_raw.get("figure_caption_overrides", raw.get("figure_caption_overrides", {}))
    )

    main_type = str(raw.get("main_type") or "").strip()
    if not main_type:
        main_type = _default_main_type_label(source_dir)

    return {
        "path": to_display_path(source_dir),
        "meta_path": to_display_path(meta_file),
        "customized": bool(meta_file.exists()),
        "main_type": main_type,
        "receptor_order": receptor_order,
        "ligand_order": ligand_order,
        "receptor_labels": receptor_labels,
        "ligand_labels": ligand_labels,
        "receptors": [{"id": key, "label": value} for key, value in receptor_labels.items()],
        "ligands": [{"id": key, "label": value} for key, value in ligand_labels.items()],
        "figure_start_number": figure_start_number,
        "extra_sections": extra_sections,
        "figure_caption_overrides": figure_caption_overrides,
        "report": {
            "figure_start_number": figure_start_number,
            "extra_sections": extra_sections,
            "figure_caption_overrides": figure_caption_overrides,
        },
    }


def _save_source_metadata(
    source_dir: Path,
    *,
    main_type: str,
    receptor_labels: dict[str, str],
    ligand_labels: dict[str, str],
    receptor_order: list[str],
    ligand_order: list[str],
    figure_start_number: int = 1,
    extra_sections: list[dict[str, str]] | None = None,
    figure_caption_overrides: dict[str, str] | None = None,
) -> dict[str, Any]:
    source_dir = source_dir.resolve()
    meta_file = _metadata_file_for_source(source_dir)
    normalized_figure_start = _normalize_positive_int(figure_start_number, default=1, min_value=1, max_value=999)
    normalized_extra_sections = _normalize_extra_sections(extra_sections or [])
    normalized_caption_overrides = _normalize_caption_map(figure_caption_overrides or {})
    payload = {
        "version": 1,
        "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        "main_type": str(main_type or "").strip(),
        "receptor_order": [str(item).strip() for item in receptor_order if str(item).strip()],
        "ligand_order": [str(item).strip() for item in ligand_order if str(item).strip()],
        "receptor_labels": {str(k): str(v).strip() for k, v in receptor_labels.items() if str(k).strip() and str(v).strip()},
        "ligand_labels": {str(k): str(v).strip() for k, v in ligand_labels.items() if str(k).strip() and str(v).strip()},
        "report": {
            "figure_start_number": normalized_figure_start,
            "extra_sections": normalized_extra_sections,
            "figure_caption_overrides": normalized_caption_overrides,
        },
    }
    meta_file.write_text(json.dumps(payload, ensure_ascii=True, indent=2, sort_keys=True), encoding="utf-8")
    receptors, ligands = _collect_source_entities(source_dir)
    return _load_source_metadata(source_dir, receptors, ligands)


def _apply_source_metadata_to_rows(rows: list[dict[str, Any]], metadata: dict[str, Any]) -> list[dict[str, Any]]:
    receptor_labels = {str(k): str(v) for k, v in (metadata.get("receptor_labels") or {}).items()}
    ligand_labels = {str(k): str(v) for k, v in (metadata.get("ligand_labels") or {}).items()}
    receptor_order = [str(item) for item in (metadata.get("receptor_order") or []) if str(item)]
    ligand_order = [str(item) for item in (metadata.get("ligand_order") or []) if str(item)]
    receptor_order_index = {name: idx for idx, name in enumerate(receptor_order)}
    ligand_order_index = {name: idx for idx, name in enumerate(ligand_order)}
    enriched: list[dict[str, Any]] = []

    for row in rows:
        row_copy = dict(row)
        receptor_id = str(row_copy.get("id") or "")
        row_copy["display_id"] = receptor_labels.get(receptor_id, _prettify_label(receptor_id))
        raw_valid_ligands = [str(item) for item in (row_copy.get("valid_ligands") or []) if str(item)]
        raw_valid_ligands.sort(key=lambda name: (ligand_order_index.get(name, 10**6), _ligand_sort_key(name)))
        row_copy["valid_ligands"] = raw_valid_ligands
        row_copy["valid_ligands_display"] = [
            ligand_labels.get(ligand_name, _prettify_label(ligand_name, trim_run_suffix=True))
            for ligand_name in raw_valid_ligands
        ]
        ligand_rows: list[dict[str, Any]] = []
        for ligand_row in row_copy.get("ligands") or []:
            ligand_item = dict(ligand_row)
            ligand_name = str(ligand_item.get("ligand") or "")
            ligand_item["display_ligand"] = ligand_labels.get(
                ligand_name,
                _prettify_label(ligand_name, trim_run_suffix=True),
            )
            ligand_rows.append(ligand_item)
        ligand_rows.sort(
            key=lambda item: (
                ligand_order_index.get(str(item.get("ligand") or ""), 10**6),
                _ligand_sort_key(str(item.get("ligand") or "")),
            )
        )
        row_copy["ligands"] = ligand_rows
        enriched.append(row_copy)
    enriched.sort(
        key=lambda item: (
            receptor_order_index.get(str(item.get("id") or ""), 10**6),
            _receptor_sort_key(str(item.get("id") or "")),
        )
    )
    return enriched


def _collect_entities_from_rows(rows: list[dict[str, Any]]) -> tuple[list[str], list[str]]:
    receptor_ids = sorted(
        [str(row.get("id") or "").strip() for row in rows if str(row.get("id") or "").strip()],
        key=_receptor_sort_key,
    )
    ligand_names: set[str] = set()
    for row in rows:
        for item in row.get("ligands") or []:
            ligand_name = str(item.get("ligand") or "").strip()
            if ligand_name:
                ligand_names.add(ligand_name)
        for ligand_name in row.get("valid_ligands") or []:
            text = str(ligand_name or "").strip()
            if text:
                ligand_names.add(text)
    return receptor_ids, sorted(ligand_names, key=_ligand_sort_key)


def _resolve_report_root(root_path: str) -> Path:
    return resolve_dock_directory(root_path, default=DOCK_DIR_RESOLVED, allow_create=False)


def resolve_dock_directory(path_text: str, *, default: Path, allow_create: bool) -> Path:
    def _rebase_to_dock(raw_text: str) -> Path | None:
        raw_candidate = Path(str(raw_text or "").strip().replace("\\", "/")).expanduser()
        parts = [part for part in raw_candidate.parts if part not in {"", "."}]
        if not parts:
            return None
        lowered = [part.lower() for part in parts]
        idx = None
        for i in range(len(lowered) - 1):
            if lowered[i] == "data" and lowered[i + 1] == "dock":
                idx = i
                break
        if idx is None:
            return None
        tail = [part for part in parts[idx + 2 :] if part not in {"", "."}]
        if any(part == ".." for part in tail):
            return None
        rebased = DOCK_DIR_RESOLVED / Path(*tail) if tail else DOCK_DIR_RESOLVED
        return rebased.resolve()

    raw = str(path_text or "").strip()
    if not raw:
        return default.resolve()
    candidate = Path(raw).expanduser()
    if not candidate.is_absolute():
        # Try WORKSPACE_DIR first (data/dock lives there), then BASE
        ws = (WORKSPACE_DIR / candidate).resolve()
        if ws.exists() and ws.is_dir():
            candidate = ws
        else:
            candidate = (BASE / candidate).resolve()
    else:
        candidate = candidate.resolve()
    if candidate != DOCK_DIR_RESOLVED and DOCK_DIR_RESOLVED not in candidate.parents:
        rebased = _rebase_to_dock(raw)
        if rebased is not None:
            candidate = rebased
        if candidate != DOCK_DIR_RESOLVED and DOCK_DIR_RESOLVED not in candidate.parents:
            raise HTTPException(status_code=400, detail="Path must be inside data/dock.")
    if candidate.exists():
        if not candidate.is_dir():
            raise HTTPException(status_code=400, detail="Path is not a directory.")
    else:
        if not allow_create:
            raise HTTPException(status_code=400, detail="Path not found.")
        candidate.mkdir(parents=True, exist_ok=True)
    return candidate


def relative_to_base(path: Path) -> str | None:
    resolved = path.resolve()
    # Accept paths inside BASE or WORKSPACE_DIR
    if (
        resolved != BASE_RESOLVED and BASE_RESOLVED not in resolved.parents
        and resolved != WORKSPACE_RESOLVED and WORKSPACE_RESOLVED not in resolved.parents
    ):
        return None
    return to_display_path(resolved)


def _default_report_source(report_root: Path) -> Path:
    report_root = report_root.resolve()
    candidates = [
        report_root / "dimer_final_linked",
        report_root / "dimer_final",
        report_root / "dimer_full",
        report_root,
    ]

    for cand in candidates:
        if not (cand.exists() and cand.is_dir()):
            continue
        rows = _collect_receptor_rows(cand)
        if any(bool(row.get("ready")) for row in rows):
            return cand.resolve()

    for cand in candidates:
        if cand.exists() and cand.is_dir():
            return cand.resolve()
    return report_root


def _find_ready_report_source(report_root: Path, current_source: Path) -> Path:
    report_root = report_root.resolve()
    current_source = current_source.resolve()
    candidates: list[Path] = [current_source, _default_report_source(report_root)]
    for directory in sorted((p for p in report_root.iterdir() if p.is_dir()), key=lambda p: p.name.lower()):
        if directory.name.startswith("."):
            continue
        candidates.append(directory.resolve())
    seen: set[Path] = set()
    for cand in candidates:
        if cand in seen:
            continue
        seen.add(cand)
        if not cand.exists() or not cand.is_dir():
            continue
        rows = _collect_receptor_rows(cand)
        if any(bool(row.get("ready")) for row in rows):
            return cand
    return current_source


def _resolve_report_source(report_root: Path, source_path: str) -> Path:
    default_source = _default_report_source(report_root)
    source_dir = resolve_dock_directory(source_path, default=default_source, allow_create=False)
    report_root = report_root.resolve()
    if source_dir != report_root and report_root not in source_dir.parents:
        raise HTTPException(status_code=400, detail="Selected source must be inside report root.")
    return source_dir


def _resolve_report_output_root(report_root: Path, source_dir: Path, output_path: str) -> Path:
    default_output = (source_dir / "report_outputs").resolve()
    output_root = resolve_dock_directory(output_path, default=default_output, allow_create=True)
    report_root = report_root.resolve()
    if output_root != report_root and report_root not in output_root.parents:
        raise HTTPException(status_code=400, detail="Output path must be inside report root.")
    return output_root


def _report_output_paths(output_root: Path) -> tuple[Path, Path, Path, Path]:
    output_root = output_root.resolve()
    render_dir = output_root / "render_images"
    plot_dir = output_root / "plots"
    doc_path = output_root / "reports" / "docking_report_mvp.docx"
    return output_root, render_dir, plot_dir, doc_path


def _sanitize_png_stem(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_.-]+", "_", str(name or "").strip()).strip("._")
    return cleaned or "figure"


def timestamp_token() -> str:
    return time.strftime("%Y%m%d_%H%M%S")


def _next_unique_png_path(directory: Path, stem: str) -> Path:
    safe_stem = _sanitize_png_stem(stem)
    candidate = directory / f"{safe_stem}.png"
    if not candidate.exists():
        return candidate
    idx = 2
    while True:
        numbered = directory / f"{safe_stem}_{idx:02d}.png"
        if not numbered.exists():
            return numbered
        idx += 1


def _is_report_image_file(path: Path) -> bool:
    return path.suffix.lower() in REPORT_IMAGE_EXTENSIONS


def _list_images_under_root(image_root: Path, *, category: str = "report") -> list[dict[str, Any]]:
    image_root = image_root.resolve()
    if not image_root.exists() or not image_root.is_dir():
        return []

    rows: list[dict[str, Any]] = []
    skip_dirs = {"__pycache__", ".tmp_render"}

    for current, dirs, files in os.walk(image_root):
        current_dir = Path(current)
        dirs[:] = [name for name in dirs if not name.startswith(".") and name not in skip_dirs]
        for filename in files:
            path = current_dir / filename
            if not _is_report_image_file(path):
                continue
            rel_path = relative_to_base(path)
            if not rel_path:
                continue
            try:
                root_rel = str(path.relative_to(image_root)).replace("\\", "/")
            except ValueError:
                root_rel = path.name
            try:
                stat = path.stat()
                mtime = int(stat.st_mtime)
                size_bytes = int(stat.st_size)
            except OSError:
                mtime = 0
                size_bytes = 0
            rows.append(
                {
                    "name": path.name,
                    "path": rel_path,
                    "root_relative_path": root_rel,
                    "category": category,
                    "kind": "image",
                    "mtime": mtime,
                    "size_bytes": size_bytes,
                }
            )

    rows.sort(key=lambda item: (-(item.get("mtime") or 0), item.get("root_relative_path") or item.get("name") or ""))
    return rows


def _resolve_report_images_root(report_root: Path, output_root: Path, images_root_path: str) -> Path:
    default_root = output_root.resolve()
    image_root = resolve_dock_directory(images_root_path, default=default_root, allow_create=False)
    report_root = report_root.resolve()
    if image_root != report_root and report_root not in image_root.parents:
        raise HTTPException(status_code=400, detail="Images root must be inside report root.")
    return image_root


def _resolve_report_image_path(report_root: Path, image_root: Path, image_path: str) -> Path:
    raw = str(image_path or "").strip()
    if not raw:
        raise HTTPException(status_code=400, detail="Missing image path.")

    candidate = Path(raw).expanduser()
    if not candidate.is_absolute():
        ws = (WORKSPACE_DIR / candidate).resolve()
        if ws.exists():
            candidate = ws
        else:
            candidate = (BASE / candidate).resolve()
    else:
        candidate = candidate.resolve()

    report_root = report_root.resolve()
    image_root = image_root.resolve()
    if candidate != report_root and report_root not in candidate.parents:
        raise HTTPException(status_code=400, detail="Image path must be inside report root.")
    if candidate != image_root and image_root not in candidate.parents:
        raise HTTPException(status_code=400, detail="Image path must be inside images root.")
    if not _is_report_image_file(candidate):
        raise HTTPException(status_code=400, detail="Unsupported image type.")
    return candidate


def _list_generated_images(directory: Path, *, category: str, kind: str) -> list[dict[str, Any]]:
    if not directory.exists() or not directory.is_dir():
        return []
    rows: list[dict[str, Any]] = []
    for png in sorted(directory.glob("*.png")):
        rel_path = relative_to_base(png)
        if not rel_path:
            continue
        try:
            stat = png.stat()
            mtime = int(stat.st_mtime)
            size_bytes = int(stat.st_size)
        except OSError:
            mtime = 0
            size_bytes = 0
        metadata = _read_image_metadata(png)
        elapsed_seconds_raw = metadata.get("elapsed_seconds")
        elapsed_seconds = None
        if isinstance(elapsed_seconds_raw, (int, float)):
            elapsed_seconds = round(max(0.0, float(elapsed_seconds_raw)), 3)
        render_dpi_raw = metadata.get("render_dpi")
        render_dpi = None
        if isinstance(render_dpi_raw, (int, float)):
            render_dpi = max(1, int(round(float(render_dpi_raw))))
        rows.append(
            {
                "name": png.name,
                "path": rel_path,
                "category": category,
                "kind": kind,
                "mtime": mtime,
                "size_bytes": size_bytes,
                "elapsed_seconds": elapsed_seconds,
                "render_dpi": render_dpi,
            }
        )
    rows.sort(key=lambda item: (-(item.get("mtime") or 0), item.get("name") or ""))
    return rows


def _collect_report_outputs(output_root: Path) -> dict[str, Any]:
    _, render_dir, plot_dir, doc_path = _report_output_paths(output_root)
    render_images = _list_generated_images(render_dir, category="render", kind="rendered")
    plot_images = _list_generated_images(plot_dir, category="plot", kind="plot")
    doc_rel = relative_to_base(doc_path) if doc_path.exists() else None
    summary = {
        "total": len(render_images) + len(plot_images),
        "rendered": len(render_images),
        "plots": len(plot_images),
        "report_ready": bool(doc_rel),
    }
    return {
        "render_images": render_images,
        "plot_images": plot_images,
        "summary": summary,
        "report_doc": {
            "exists": bool(doc_rel),
            "path": doc_rel,
        },
    }


def _run_sort_key(name: str) -> tuple[int, int, str]:
    match = REPORT_RUN_RE.fullmatch(name or "")
    if match:
        return (0, int(match.group(1)), name.lower())
    return (1, 10**9, (name or "").lower())


def _is_receptor_dtype(name: str) -> bool:
    return bool(REPORT_DTYPE_RE.fullmatch(name or ""))


def _is_valid_run_payload_dir(run_dir: Path) -> bool:
    if not run_dir.exists() or not run_dir.is_dir():
        return False
    if not any(run_dir.glob("*_complex.pdb")):
        return False
    if not (run_dir / "interaction_map.json").exists():
        return False
    if not (run_dir / "plip" / "report.xml").exists():
        return False
    return True


def _valid_run_dirs(ligand_dir: Path) -> list[Path]:
    if not ligand_dir.exists() or not ligand_dir.is_dir():
        return []
    valid: list[Path] = []
    for run_dir in (p for p in ligand_dir.iterdir() if p.is_dir()):
        if _is_valid_run_payload_dir(run_dir):
            valid.append(run_dir)
    valid.sort(key=lambda p: _run_sort_key(p.name))
    return valid


def _count_valid_runs(ligand_dir: Path) -> int:
    return len(_valid_run_dirs(ligand_dir))


def _candidate_run_dirs(ligand_dir: Path) -> list[tuple[str, Path]]:
    if not ligand_dir.exists() or not ligand_dir.is_dir():
        return []
    entries: list[tuple[str, Path]] = []
    for run_dir in (p for p in ligand_dir.iterdir() if p.is_dir()):
        if not REPORT_RUN_RE.fullmatch(run_dir.name):
            continue
        entries.append((run_dir.name, run_dir.resolve()))
    entries.sort(key=lambda item: _run_sort_key(item[0]))
    return entries


def _collect_receptor_inventory(source_dir: Path) -> dict[str, dict[str, list[tuple[str, Path]]]]:
    source_dir = source_dir.resolve()
    inventory: dict[str, dict[str, list[tuple[str, Path]]]] = {}
    if not source_dir.exists() or not source_dir.is_dir():
        return inventory

    skip_names = {"report_outputs", "plots", "render_images", "reports", "__pycache__", ".tmp_render"}

    def _add_run(receptor_id: str, ligand_name: str, run_name: str, run_dir: Path) -> None:
        receptor_key = str(receptor_id or "").strip()
        ligand_key = str(ligand_name or "").strip()
        run_key = str(run_name or "").strip()
        if not receptor_key or not ligand_key or not run_key:
            return
        receptor_bucket = inventory.setdefault(receptor_key, {})
        run_bucket = receptor_bucket.setdefault(ligand_key, [])
        resolved = run_dir.resolve()
        if any((entry_name == run_key and entry_dir == resolved) for entry_name, entry_dir in run_bucket):
            return
        run_bucket.append((run_key, resolved))

    def _collect_from_receptor_dir(receptor_dir: Path, receptor_id: str) -> None:
        if not receptor_dir.exists() or not receptor_dir.is_dir():
            return
        for ligand_dir in sorted((p for p in receptor_dir.iterdir() if p.is_dir()), key=lambda p: p.name.lower()):
            if ligand_dir.name.startswith(".") or ligand_dir.name in skip_names:
                continue
            run_dirs = _valid_run_dirs(ligand_dir)
            if not run_dirs:
                continue
            # Hierarchical receptor layout expects run directories named runN.
            # This avoids treating report roots (e.g. dimer_full) as receptor/ligand rows.
            if not any(REPORT_RUN_RE.fullmatch(run_dir.name) for run_dir in run_dirs):
                continue
            for run_dir in run_dirs:
                _add_run(receptor_id, ligand_dir.name, run_dir.name, run_dir)

    # Pattern A/B: source/receptor/ligand/runX and source-as-receptor/ligand/runX
    _collect_from_receptor_dir(source_dir, source_dir.name)
    for receptor_dir in sorted((p for p in source_dir.iterdir() if p.is_dir()), key=lambda p: p.name.lower()):
        if receptor_dir.name.startswith(".") or receptor_dir.name in skip_names:
            continue
        _collect_from_receptor_dir(receptor_dir, receptor_dir.name)

    # Pattern C: source/PDB_ligand_runN
    for case_dir in sorted((p for p in source_dir.iterdir() if p.is_dir()), key=lambda p: p.name.lower()):
        if case_dir.name.startswith(".") or case_dir.name in skip_names:
            continue
        if not _is_valid_run_payload_dir(case_dir):
            continue
        match = REPORT_CASE_RUN_RE.fullmatch(case_dir.name)
        if not match:
            continue
        receptor_id = match.group(1)
        ligand_name = match.group(2)
        run_name = f"run{int(match.group(3))}"
        _add_run(receptor_id, ligand_name, run_name, case_dir)

    for receptor_id, ligand_map in inventory.items():
        for ligand_name, run_entries in ligand_map.items():
            run_entries.sort(key=lambda item: _run_sort_key(item[0]))

    return inventory


def _collect_receptor_candidates(source_dir: Path) -> dict[str, dict[str, set[str]]]:
    source_dir = source_dir.resolve()
    candidates: dict[str, dict[str, set[str]]] = {}
    if not source_dir.exists() or not source_dir.is_dir():
        return candidates

    skip_names = {"report_outputs", "plots", "render_images", "reports", "__pycache__", ".tmp_render"}

    def _add(receptor_id: str, ligand_name: str, run_name: str) -> None:
        receptor_key = str(receptor_id or "").strip()
        ligand_key = str(ligand_name or "").strip()
        run_key = str(run_name or "").strip()
        if not receptor_key or not ligand_key:
            return
        receptor_bucket = candidates.setdefault(receptor_key, {})
        run_bucket = receptor_bucket.setdefault(ligand_key, set())
        if run_key:
            run_bucket.add(run_key)

    def _collect_from_receptor_dir(receptor_dir: Path, receptor_id: str) -> None:
        if not receptor_dir.exists() or not receptor_dir.is_dir():
            return
        for ligand_dir in sorted((p for p in receptor_dir.iterdir() if p.is_dir()), key=lambda p: p.name.lower()):
            if ligand_dir.name.startswith(".") or ligand_dir.name in skip_names:
                continue
            run_entries = _candidate_run_dirs(ligand_dir)
            if not run_entries:
                continue
            for run_name, _ in run_entries:
                _add(receptor_id, ligand_dir.name, run_name)

    _collect_from_receptor_dir(source_dir, source_dir.name)
    for receptor_dir in sorted((p for p in source_dir.iterdir() if p.is_dir()), key=lambda p: p.name.lower()):
        if receptor_dir.name.startswith(".") or receptor_dir.name in skip_names:
            continue
        _collect_from_receptor_dir(receptor_dir, receptor_dir.name)

    for case_dir in sorted((p for p in source_dir.iterdir() if p.is_dir()), key=lambda p: p.name.lower()):
        if case_dir.name.startswith(".") or case_dir.name in skip_names:
            continue
        match = REPORT_CASE_RUN_RE.fullmatch(case_dir.name)
        if not match:
            continue
        receptor_id = match.group(1)
        ligand_name = match.group(2)
        run_name = f"run{int(match.group(3))}"
        _add(receptor_id, ligand_name, run_name)

    return candidates


def _looks_like_linked_root(directory: Path) -> bool:
    if not directory.exists() or not directory.is_dir():
        return False
    receptor_dirs = [p for p in directory.iterdir() if p.is_dir() and _is_receptor_dtype(p.name)]
    if not receptor_dirs:
        return False
    for receptor_dir in receptor_dirs:
        for ligand_dir in (p for p in receptor_dir.iterdir() if p.is_dir()):
            if ligand_dir.name.endswith("_1"):
                return True
    return False


def _discover_linked_roots(source_dir: Path, max_depth: int = REPORT_MAX_DISCOVERY_DEPTH) -> list[Path]:
    source_dir = source_dir.resolve()
    if not source_dir.exists() or not source_dir.is_dir():
        return []

    skip_names = {
        "__pycache__",
        "plip",
        "plots",
        "render_images",
        "reports",
        "report_outputs",
        ".tmp_render",
    }
    seen: set[Path] = set()
    found: list[Path] = []

    def _add_candidate(path: Path) -> None:
        resolved = path.resolve()
        if resolved in seen:
            return
        if _looks_like_linked_root(resolved):
            seen.add(resolved)
            found.append(resolved)

    _add_candidate(source_dir)

    for current, dirs, _ in os.walk(source_dir):
        current_dir = Path(current)
        try:
            depth = len(current_dir.relative_to(source_dir).parts)
        except ValueError:
            continue

        if depth >= max_depth:
            dirs[:] = []
            continue

        pruned: list[str] = []
        for name in dirs:
            lower = name.lower()
            if name.startswith("."):
                continue
            if name in skip_names or lower in skip_names:
                continue
            if REPORT_RUN_RE.fullmatch(name):
                continue
            if lower.endswith("_results"):
                continue
            if "_run" in lower:
                continue
            pruned.append(name)
        dirs[:] = pruned

        for child_name in list(dirs):
            child_dir = current_dir / child_name
            if _looks_like_linked_root(child_dir):
                _add_candidate(child_dir)
                dirs.remove(child_name)

    found.sort(
        key=lambda path: (
            len(path.relative_to(source_dir).parts) if path == source_dir or source_dir in path.parents else 10**6,
            path.name.lower(),
            str(path).lower(),
        )
    )
    return found


def _resolve_dimer_linked_root(source_dir: Path, _report_root: Path, linked_path: str = "") -> tuple[Path, list[Path]]:
    source_dir = source_dir.resolve()

    candidates: list[Path] = _discover_linked_roots(source_dir)

    unique_candidates: list[Path] = []
    seen: set[Path] = set()
    for candidate in candidates:
        resolved = candidate.resolve()
        if resolved in seen:
            continue
        seen.add(resolved)
        unique_candidates.append(resolved)

    if not unique_candidates:
        raise HTTPException(
            status_code=400,
            detail="No render-ready docking roots found under selected source.",
        )

    requested: Path | None = None
    if linked_path:
        try:
            requested = resolve_dock_directory(linked_path, default=source_dir, allow_create=False).resolve()
        except HTTPException:
            requested = None

    selected: Path | None = None
    if requested is not None:
        for candidate in unique_candidates:
            if requested == candidate:
                selected = candidate
                break
        if selected is None:
            for candidate in unique_candidates:
                if candidate in requested.parents:
                    selected = candidate
                    break
        if selected is None:
            for candidate in unique_candidates:
                if requested in candidate.parents:
                    selected = candidate
                    break

    if selected is None:
        def _rank(path: Path) -> tuple[int, int, int, str]:
            if path == source_dir or source_dir in path.parents:
                scope_rank = 0
                depth = len(path.relative_to(source_dir).parts) if path != source_dir else 0
            else:
                scope_rank = 1
                depth = 10**6
            linked_name_rank = 0 if path.name.lower() == "dimer_final_linked" else 1
            return (scope_rank, depth, linked_name_rank, str(path).lower())

        selected = sorted(unique_candidates, key=_rank)[0]

    return selected, unique_candidates


def _find_render_inputs(
    inventory: dict[str, dict[str, list[tuple[str, Path]]]],
    receptor_id: str,
    ligand_name: str,
    preferred_run: str = "",
) -> tuple[Path, Path, Path, str]:
    ligand_runs = (inventory.get(receptor_id) or {}).get(ligand_name) or []
    if not ligand_runs:
        raise FileNotFoundError(f"No valid run entries for {receptor_id}/{ligand_name}")

    run_entry: tuple[str, Path] | None = None
    if preferred_run:
        for run_name, run_dir in ligand_runs:
            if run_name == preferred_run:
                run_entry = (run_name, run_dir)
                break
    if run_entry is None:
        for run_name, run_dir in ligand_runs:
            if run_name.lower() == "run1":
                run_entry = (run_name, run_dir)
                break
    if run_entry is None:
        run_entry = sorted(ligand_runs, key=lambda item: _run_sort_key(item[0]))[-1]

    selected_run_name, selected_run_dir = run_entry
    complex_files = sorted(selected_run_dir.glob("*_complex.pdb"))
    if not complex_files:
        raise FileNotFoundError(
            f"Missing complex pdb for {receptor_id}/{ligand_name} in {selected_run_name}"
        )
    interaction_json = selected_run_dir / "interaction_map.json"
    plip_report = selected_run_dir / "plip" / "report.xml"
    return complex_files[0], interaction_json, plip_report, selected_run_name


def _select_otofigure_ligand_runs(
    inventory: dict[str, dict[str, list[tuple[str, Path]]]],
    receptor_id: str,
    *,
    preferred_ligand: str = "",
    ligand_order_index: dict[str, int] | None = None,
) -> tuple[str, list[tuple[str, Path]]]:
    ligand_map = inventory.get(receptor_id) or {}
    if not ligand_map:
        raise FileNotFoundError(f"No ligand/run data found for receptor: {receptor_id}")

    preferred = str(preferred_ligand or "").strip()
    if preferred:
        run_entries = sorted(ligand_map.get(preferred) or [], key=lambda item: _run_sort_key(item[0]))
        if not run_entries:
            raise FileNotFoundError(f"No multi-run ligand data found for receptor {receptor_id}: {preferred}")
        return preferred, run_entries[:5]

    order_index = dict(ligand_order_index or {})
    ranked: list[tuple[int, int, tuple[str, str], str, list[tuple[str, Path]]]] = []
    for ligand_name, run_entries in ligand_map.items():
        valid_runs = sorted(run_entries, key=lambda item: _run_sort_key(item[0]))
        if not valid_runs:
            continue
        ranked.append(
            (
                -len(valid_runs),
                order_index.get(ligand_name, 10**6),
                _ligand_sort_key(ligand_name),
                ligand_name,
                valid_runs,
            )
        )

    if not ranked:
        raise FileNotFoundError(f"No multi-run ligand data found for receptor: {receptor_id}")

    ranked.sort()
    _, _, _, ligand_name, run_entries = ranked[0]
    return ligand_name, run_entries[:5]


def _collect_receptor_rows(source_dir: Path) -> list[dict[str, Any]]:
    inventory = _collect_receptor_inventory(source_dir)
    candidates = _collect_receptor_candidates(source_dir)
    rows: list[dict[str, Any]] = []

    receptor_ids = set(inventory.keys()) | set(candidates.keys())

    for receptor_id in sorted(receptor_ids, key=_receptor_sort_key):
        valid_ligand_map = inventory.get(receptor_id, {})
        candidate_ligand_map = candidates.get(receptor_id, {})
        ligand_status: list[dict[str, Any]] = []
        valid_run_counts: list[int] = []
        run_union: set[str] = set()

        ligand_names = sorted(set(valid_ligand_map.keys()) | set(candidate_ligand_map.keys()), key=_ligand_sort_key)
        for ligand_name in ligand_names:
            run_entries = valid_ligand_map.get(ligand_name, [])
            valid_run_names = [run_name for run_name, _ in run_entries]
            candidate_run_names = sorted(candidate_ligand_map.get(ligand_name, set()), key=_run_sort_key)
            run_names = sorted(set(valid_run_names) | set(candidate_run_names), key=_run_sort_key)
            valid_runs = len(valid_run_names)
            if valid_runs > 0:
                valid_run_counts.append(valid_runs)
                run_union.update(valid_run_names)
            ligand_status.append(
                {
                    "ligand": ligand_name,
                    "valid_runs": valid_runs,
                    "detected_runs": len(run_names),
                    "runs": run_names,
                }
            )

        run_options = sorted(run_union, key=_run_sort_key)
        default_run = "run1" if "run1" in run_options else (run_options[0] if run_options else "")
        ready = bool(valid_run_counts)

        rows.append(
            {
                "id": receptor_id,
                "ready": ready,
                "valid_ligands": [item["ligand"] for item in ligand_status if item["valid_runs"] > 0] or [item["ligand"] for item in ligand_status],
                "runs_per_ligand_min": min(valid_run_counts) if valid_run_counts else 0,
                "total_valid_runs": sum(valid_run_counts),
                "run_options": run_options,
                "default_run": default_run,
                "ligands": ligand_status,
            }
        )

    return rows


def _describe_linked_root(linked_root: Path, *, selected: bool = False) -> dict[str, Any]:
    receptor_rows = _collect_receptor_rows(linked_root)
    ready_rows = [row for row in receptor_rows if row.get("ready")]
    docking_count = sum(int(row.get("runs_per_ligand_min") or 0) for row in ready_rows)
    return {
        "name": linked_root.name,
        "path": to_display_path(linked_root),
        "selected": selected,
        "receptor_count": len(receptor_rows),
        "ready_receptors": len(ready_rows),
        "docking_count": docking_count,
    }


def _list_source_folders(report_root: Path, selected_source: Path | None = None) -> list[dict[str, Any]]:
    selected_source_resolved = selected_source.resolve() if selected_source else None
    # Skip internal/system directories that are not valid report sources
    SKIP_NAMES = {
        "report_outputs", "__pycache__", "_run_sessions", "_meta",
        "plip", "plots", "render_images", "reports", ".tmp_render",
    }
    entries: list[dict[str, Any]] = []
    for directory in sorted((p for p in report_root.iterdir() if p.is_dir()), key=lambda p: p.name.lower()):
        if directory.name.startswith((".", "_")):
            continue
        if directory.name in SKIP_NAMES:
            continue
        receptor_rows = _collect_receptor_rows(directory)
        receptor_ids, ligand_names = _collect_entities_from_rows(receptor_rows)
        source_meta = _load_source_metadata(directory, receptor_ids, ligand_names)
        ready_rows = [row for row in receptor_rows if row.get("ready")]
        docking_count = sum(int(row.get("runs_per_ligand_min") or 0) for row in ready_rows)
        entries.append(
            {
                "name": directory.name,
                "path": to_display_path(directory.resolve()),
                "selected": bool(selected_source_resolved and directory.resolve() == selected_source_resolved),
                "receptor_count": len(receptor_rows),
                "ready_receptors": len(ready_rows),
                "docking_count": docking_count,
                "main_type": source_meta.get("main_type", ""),
            }
        )
    return entries


def _load_interaction_payload(interaction_json: Path) -> dict[str, Any]:
    try:
        payload = json.loads(interaction_json.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return payload if isinstance(payload, dict) else {}


def _interaction_type_palette() -> dict[str, str]:
    return {
        "hydrophobic": "#0f766e",
        "hbond": "#2563eb",
        "hydrogen bond": "#2563eb",
        "pi-stacking": "#7c3aed",
        "pi_cation": "#9333ea",
        "electrostatic": "#dc2626",
        "salt bridge": "#dc2626",
        "halogen": "#ea580c",
        "water bridge": "#0891b2",
    }


def _build_interaction_preview_svg(
    *,
    receptor_label: str,
    ligand_label: str,
    run_name: str,
    render_mode: str,
    contact_count: int,
    residue_rows: list[dict[str, Any]],
    type_counts: dict[str, int],
) -> str:
    width = 720
    height = 360
    left = 36
    right = width - 36
    chart_width = 276
    top_residues = residue_rows[:6]
    type_rows = sorted(type_counts.items(), key=lambda item: (-item[1], item[0].lower()))[:6]
    max_residue = max((int(item.get("contact_count") or 0) for item in top_residues), default=1)
    max_type = max((int(count or 0) for _, count in type_rows), default=1)
    palette = _interaction_type_palette()

    parts: list[str] = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}" fill="none">',
        '<defs>',
        '<linearGradient id="previewBg" x1="0" y1="0" x2="1" y2="1">',
        '<stop offset="0%" stop-color="#f8fafc"/>',
        '<stop offset="100%" stop-color="#eef6ff"/>',
        '</linearGradient>',
        '</defs>',
        f'<rect x="0" y="0" width="{width}" height="{height}" rx="18" fill="url(#previewBg)"/>',
        f'<rect x="18" y="18" width="{width - 36}" height="{height - 36}" rx="16" fill="#ffffff" stroke="#dbe3ef"/>',
        f'<text x="{left}" y="52" font-size="13" font-family="IBM Plex Sans, Arial, sans-serif" fill="#64748b">Panel Preview</text>',
        f'<text x="{left}" y="78" font-size="24" font-weight="700" font-family="IBM Plex Sans, Arial, sans-serif" fill="#0f172a">{html.escape(receptor_label)} · {html.escape(ligand_label)}</text>',
        f'<text x="{left}" y="102" font-size="12" font-family="IBM Plex Sans, Arial, sans-serif" fill="#475569">Run {html.escape(run_name)} · {html.escape(render_mode.title())} mode · {contact_count} contacts</text>',
        f'<text x="{left}" y="138" font-size="12" font-family="IBM Plex Sans, Arial, sans-serif" fill="#64748b">Top interacting residues</text>',
        f'<text x="{left + 344}" y="138" font-size="12" font-family="IBM Plex Sans, Arial, sans-serif" fill="#64748b">Interaction classes</text>',
    ]

    if not top_residues:
        parts.append(
            f'<text x="{left}" y="190" font-size="14" font-family="IBM Plex Sans, Arial, sans-serif" fill="#94a3b8">No interaction map data available for this run.</text>'
        )
    else:
        for idx, item in enumerate(top_residues):
            y = 162 + idx * 28
            label = f'{item.get("receptor_resname", "-")}{item.get("receptor_resid", "")} {item.get("receptor_chain", "")}'.strip()
            count = int(item.get("contact_count") or 0)
            min_distance = item.get("min_distance")
            bar_width = int((count / max_residue) * chart_width) if max_residue else 0
            distance_text = f'{float(min_distance):.2f} Å' if min_distance not in (None, "") else "-"
            parts.extend(
                [
                    f'<text x="{left}" y="{y}" font-size="12" font-family="IBM Plex Mono, monospace" fill="#0f172a">{html.escape(label)}</text>',
                    f'<rect x="{left}" y="{y + 7}" width="{chart_width}" height="10" rx="5" fill="#e2e8f0"/>',
                    f'<rect x="{left}" y="{y + 7}" width="{max(12, bar_width)}" height="10" rx="5" fill="#0f766e"/>',
                    f'<text x="{left + chart_width + 10}" y="{y + 16}" font-size="11" font-family="IBM Plex Sans, Arial, sans-serif" fill="#334155">{count} contacts · {html.escape(distance_text)}</text>',
                ]
            )

    if type_rows:
        for idx, (kind, count) in enumerate(type_rows):
            y = 162 + idx * 28
            color = palette.get(kind.lower(), "#475569")
            label = kind.replace("_", " ")
            bar_width = int((count / max_type) * chart_width) if max_type else 0
            parts.extend(
                [
                    f'<text x="{left + 344}" y="{y}" font-size="12" font-family="IBM Plex Sans, Arial, sans-serif" fill="#0f172a">{html.escape(label.title())}</text>',
                    f'<rect x="{left + 344}" y="{y + 7}" width="{chart_width}" height="10" rx="5" fill="#e2e8f0"/>',
                    f'<rect x="{left + 344}" y="{y + 7}" width="{max(12, bar_width)}" height="10" rx="5" fill="{color}"/>',
                    f'<text x="{left + 344 + chart_width + 10}" y="{y + 16}" font-size="11" font-family="IBM Plex Sans, Arial, sans-serif" fill="#334155">{count}</text>',
                ]
            )
    else:
        parts.append(
            f'<text x="{left + 344}" y="190" font-size="14" font-family="IBM Plex Sans, Arial, sans-serif" fill="#94a3b8">No interaction classes recorded.</text>'
        )

    parts.extend(
        [
            f'<rect x="{left}" y="{height - 64}" width="{width - 72}" height="1" fill="#e2e8f0"/>',
            f'<text x="{left}" y="{height - 36}" font-size="12" font-family="IBM Plex Sans, Arial, sans-serif" fill="#64748b">Preview uses the currently selected receptor and run. Render output remains unchanged.</text>',
            '</svg>',
        ]
    )
    return "".join(parts)


def _resolve_preview_context(
    source_dir: Path,
    *,
    receptor_id: str,
    run_name: str,
    render_mode: str,
) -> dict[str, Any]:
    receptor_rows = _collect_receptor_rows(source_dir)
    receptor_ids, ligand_names = _collect_entities_from_rows(receptor_rows)
    source_metadata = _load_source_metadata(source_dir, receptor_ids, ligand_names)
    receptor_labels = {str(k): str(v) for k, v in (source_metadata.get("receptor_labels") or {}).items()}
    ligand_labels = {str(k): str(v) for k, v in (source_metadata.get("ligand_labels") or {}).items()}
    ligand_order_index = {
        str(name): idx
        for idx, name in enumerate(source_metadata.get("ligand_order") or [])
        if str(name)
    }

    ready_rows = [row for row in receptor_rows if row.get("ready")]
    if not ready_rows:
        raise FileNotFoundError("No render-ready receptor found for preview.")

    row_by_id = {str(row.get("id") or ""): row for row in receptor_rows}
    selected_row = row_by_id.get(str(receptor_id or "").strip()) or ready_rows[0]
    receptor_key = str(selected_row.get("id") or "").strip()

    inventory = _collect_receptor_inventory(source_dir)
    ligand_map = inventory.get(receptor_key) or {}
    if not ligand_map:
        raise FileNotFoundError(f"No valid ligand/run data found for receptor: {receptor_key}")

    normalized_mode = _normalize_render_mode(render_mode)
    available_runs: list[str] = []
    if normalized_mode == REPORT_RENDER_MODE_OTOFIGURE:
        ligand_name, run_entries = _select_otofigure_ligand_runs(
            inventory,
            receptor_key,
            ligand_order_index=ligand_order_index,
        )
        available_runs = [name for name, _ in run_entries]
    else:
        ligand_names = sorted(
            ligand_map.keys(),
            key=lambda ligand_name: (
                ligand_order_index.get(ligand_name, 10**6),
                _ligand_sort_key(ligand_name),
            ),
        )
        ligand_name = ligand_names[0]
        available_runs = [name for name, _ in (ligand_map.get(ligand_name) or [])]

    preferred_run = str(run_name or "").strip()
    complex_pdb, interaction_json, _plip_report, selected_run_name = _find_render_inputs(
        inventory,
        receptor_key,
        ligand_name,
        preferred_run=preferred_run,
    )
    interaction_payload = _load_interaction_payload(interaction_json)
    contacts = interaction_payload.get("contacts") if isinstance(interaction_payload.get("contacts"), list) else []
    residue_rows = interaction_payload.get("residue_summary") if isinstance(interaction_payload.get("residue_summary"), list) else []
    residue_rows = sorted(
        [row for row in residue_rows if isinstance(row, dict)],
        key=lambda item: (-int(item.get("contact_count") or 0), float(item.get("min_distance") or 10**6)),
    )
    type_counts: dict[str, int] = {}
    for contact in contacts:
        if not isinstance(contact, dict):
            continue
        kind = str(contact.get("interaction_type") or "other").strip().lower() or "other"
        type_counts[kind] = type_counts.get(kind, 0) + 1

    receptor_label = receptor_labels.get(receptor_key, _prettify_label(receptor_key))
    ligand_label = ligand_labels.get(ligand_name, _prettify_label(ligand_name, trim_run_suffix=True))
    svg_markup = _build_interaction_preview_svg(
        receptor_label=receptor_label,
        ligand_label=ligand_label,
        run_name=selected_run_name,
        render_mode=normalized_mode,
        contact_count=int(interaction_payload.get("contact_count") or len(contacts) or 0),
        residue_rows=residue_rows,
        type_counts=type_counts,
    )

    return {
        "available": True,
        "receptor_id": receptor_key,
        "receptor_label": receptor_label,
        "ligand_name": ligand_name,
        "ligand_label": ligand_label,
        "run_name": selected_run_name,
        "available_runs": available_runs,
        "render_mode": normalized_mode,
        "contact_count": int(interaction_payload.get("contact_count") or len(contacts) or 0),
        "top_residues": residue_rows[:6],
        "interaction_type_counts": type_counts,
        "has_interaction_map": bool(interaction_payload),
        "complex_path": to_display_path(complex_pdb.parent),
        "svg_markup": svg_markup,
    }


def _render_dtype_panel_with_runner(
    dtype: str,
    inventory: dict[str, dict[str, list[tuple[str, Path]]]],
    output_dir: Path,
    temp_root: Path,
    dpi: int,
    *,
    pipeline_runner,
    preferred_run: str = "run1",
    preferred_ligand: str = "",
    output_stem: str = "",
    preview_mode: bool = False,
    ligand_order_index: dict[str, int] | None = None,
    otofigure_style: str = "balanced",
    otofigure_ray_trace: bool = True,
    otofigure_options: dict[str, Any] | None = None,
    process_hooks: dict[str, Any] | None = None,
) -> tuple[Path, list[str]]:
    from PIL import Image

    from figure_scripts.panel_figure.config import TargetConfig
    from figure_scripts.panel_figure.panel import add_left_margin_label, compose_2x2, load_bold_serif_font
    output_dir.mkdir(parents=True, exist_ok=True)
    temp_root.mkdir(parents=True, exist_ok=True)
    tiles: dict[str, Image.Image] = {}
    used_runs: set[str] = set()
    font = load_bold_serif_font(40)
    effective_dpi = max(30, min(600, int(dpi or 120)))
    if preview_mode:
        effective_dpi = min(effective_dpi, 72)

    ligand_map = inventory.get(dtype) or {}
    if not ligand_map:
        raise FileNotFoundError(f"No ligand/run data found for receptor: {dtype}")

    order_index = dict(ligand_order_index or {})
    ligand_names = sorted(
        ligand_map.keys(),
        key=lambda ligand_name: (order_index.get(ligand_name, 10**6), _ligand_sort_key(ligand_name)),
    )
    max_tiles = 1 if preview_mode else 4
    if preferred_ligand:
        if preferred_ligand not in ligand_map:
            raise FileNotFoundError(f"No ligand data found for receptor {dtype}: {preferred_ligand}")
        selected_ligands = [preferred_ligand]
    else:
        selected_ligands = ligand_names[:max_tiles]
    if not selected_ligands:
        raise FileNotFoundError(f"No renderable ligands found for receptor: {dtype}")

    slot_labels = ("A", "B", "C", "D")
    sample_size = (640, 480)

    for tile_label, ligand_name in zip(slot_labels, selected_ligands):
        safe_ligand = re.sub(r"[^A-Za-z0-9_.-]+", "_", ligand_name)
        work_dir = temp_root / f"{dtype}_{safe_ligand}"
        if work_dir.exists():
            shutil.rmtree(work_dir, ignore_errors=True)
        work_dir.mkdir(parents=True, exist_ok=True)
        try:
            complex_pdb, interaction_json, plip_report, run_name = _find_render_inputs(
                inventory,
                dtype,
                ligand_name,
                preferred_run=preferred_run,
            )
            used_runs.add(run_name)
            cfg = TargetConfig(
                name=f"{dtype}_{safe_ligand}",
                complex_pdb=str(complex_pdb),
                interaction_json=str(interaction_json),
                plip_report_txt=str(plip_report),
                plip_contacts_csv="",
                output_dir=str(work_dir),
                dpi=effective_dpi,
                show_labels=False,
                cleanup_intermediate=True,
                contacts_zoom=0.0,
            )
            result = pipeline_runner(cfg)
            combined_path = Path(result["combined_transparent"])
            with Image.open(combined_path) as image_obj:
                tile = image_obj.convert("RGBA").copy()
            sample_size = tile.size
            tiles[tile_label] = add_left_margin_label(
                tile,
                tile_label,
                font,
                pad=18,
                stroke_width=1,
                min_left_margin=90,
            )
        finally:
            shutil.rmtree(work_dir, ignore_errors=True)

    for tile_label in slot_labels:
        if tile_label in tiles:
            continue
        blank = Image.new("RGBA", sample_size, (255, 255, 255, 0))
        tiles[tile_label] = add_left_margin_label(
            blank,
            tile_label,
            font,
            pad=18,
            stroke_width=1,
            min_left_margin=90,
        )

    panel = compose_2x2(tiles, order=slot_labels)
    panel_stem = output_stem or f"{dtype}_{timestamp_token()}"
    out_path = _next_unique_png_path(output_dir, panel_stem)
    panel.save(out_path, dpi=(effective_dpi, effective_dpi))
    return out_path, sorted(used_runs, key=_run_sort_key)


def _render_dtype_panel(
    dtype: str,
    inventory: dict[str, dict[str, list[tuple[str, Path]]]],
    output_dir: Path,
    temp_root: Path,
    dpi: int,
    preferred_run: str = "run1",
    preferred_ligand: str = "",
    output_stem: str = "",
    preview_mode: bool = False,
    ligand_order_index: dict[str, int] | None = None,
    otofigure_style: str = "balanced",
    otofigure_ray_trace: bool = True,
    otofigure_options: dict[str, Any] | None = None,
    process_hooks: dict[str, Any] | None = None,
) -> tuple[Path, list[str]]:
    from figure_scripts.panel_figure.pipeline import run as run_pipeline

    return _render_dtype_panel_with_runner(
        dtype,
        inventory,
        output_dir,
        temp_root,
        dpi,
        pipeline_runner=run_pipeline,
        preferred_run=preferred_run,
        preferred_ligand=preferred_ligand,
        output_stem=output_stem,
        preview_mode=preview_mode,
        ligand_order_index=ligand_order_index,
        otofigure_style=otofigure_style,
        otofigure_ray_trace=otofigure_ray_trace,
        otofigure_options=otofigure_options,
        process_hooks=process_hooks,
    )


def _render_dtype_otofigure_panel(
    dtype: str,
    inventory: dict[str, dict[str, list[tuple[str, Path]]]],
    output_dir: Path,
    temp_root: Path,
    dpi: int,
    preferred_run: str = "run1",
    preferred_ligand: str = "",
    output_stem: str = "",
    preview_mode: bool = False,
    ligand_order_index: dict[str, int] | None = None,
    otofigure_style: str = "balanced",
    otofigure_ray_trace: bool = True,
    otofigure_options: dict[str, Any] | None = None,
    process_hooks: dict[str, Any] | None = None,
) -> tuple[Path, list[str]]:
    from figure_scripts.otofigure.pipeline import run as run_pipeline

    output_dir.mkdir(parents=True, exist_ok=True)
    temp_root.mkdir(parents=True, exist_ok=True)

    ligand_name, run_entries = _select_otofigure_ligand_runs(
        inventory,
        dtype,
        preferred_ligand=preferred_ligand,
        ligand_order_index=ligand_order_index,
    )
    safe_ligand = re.sub(r"[^A-Za-z0-9_.-]+", "_", ligand_name)
    work_dir = temp_root / f"otofigure_{dtype}_{safe_ligand}"
    out_path = _next_unique_png_path(output_dir, output_stem or f"{dtype}_{safe_ligand}_{timestamp_token()}")

    try:
        result = run_pipeline(
            receptor_id=dtype,
            ligand_name=ligand_name,
            run_entries=run_entries,
            output_png=out_path,
            work_dir=work_dir,
            dpi=dpi,
            style_preset=otofigure_style,
            ray_trace=otofigure_ray_trace,
            options=otofigure_options,
            preview_mode=preview_mode,
            on_process_start=(process_hooks or {}).get("on_process_start"),
            on_process_end=(process_hooks or {}).get("on_process_end"),
        )
        used_runs = [str(name) for name in result.get("used_runs") or [run_name for run_name, _ in run_entries]]
        return out_path, sorted(used_runs, key=_run_sort_key)
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


def _render_dtype_multi_ligand_panel(
    dtype: str,
    inventory: dict[str, dict[str, list[tuple[str, Path]]]],
    output_dir: Path,
    temp_root: Path,
    dpi: int,
    preferred_run: str = "run1",
    preferred_ligand: str = "",
    output_stem: str = "",
    preview_mode: bool = False,
    ligand_order_index: dict[str, int] | None = None,
    otofigure_style: str = "balanced",
    otofigure_ray_trace: bool = True,
    otofigure_options: dict[str, Any] | None = None,
    process_hooks: dict[str, Any] | None = None,
) -> tuple[Path, list[str]]:
    from figure_scripts.otofigure.multi_ligand_pipeline import run as run_pipeline

    output_dir.mkdir(parents=True, exist_ok=True)
    temp_root.mkdir(parents=True, exist_ok=True)

    ligand_map = inventory.get(dtype) or {}
    if not ligand_map:
        raise FileNotFoundError(f"No ligand/run data found for receptor: {dtype}")

    order_index = dict(ligand_order_index or {})
    ligand_names = sorted(
        ligand_map.keys(),
        key=lambda ligand_name: (order_index.get(ligand_name, 10**6), _ligand_sort_key(ligand_name)),
    )
    if preferred_ligand:
        if preferred_ligand not in ligand_map:
            raise FileNotFoundError(f"No ligand data found for receptor {dtype}: {preferred_ligand}")
        ligand_names = [preferred_ligand]
    if not ligand_names:
        raise FileNotFoundError(f"No ligand data found for receptor: {dtype}")

    selected_ligand = ligand_names[0]
    run_entries = sorted(ligand_map.get(selected_ligand) or [], key=lambda item: _run_sort_key(item[0]))
    if not run_entries:
        raise FileNotFoundError(f"No runs found for receptor {dtype}: {selected_ligand}")
    selected_run_name, selected_run_dir = run_entries[0]
    if preferred_run:
        for run_name, run_dir in run_entries:
            if run_name == preferred_run:
                selected_run_name, selected_run_dir = run_name, run_dir
                break
    elif any(run_name == "run1" for run_name, _ in run_entries):
        for run_name, run_dir in run_entries:
            if run_name == "run1":
                selected_run_name, selected_run_dir = run_name, run_dir
                break

    if not (selected_run_dir / "multi_ligand" / "sites.json").exists():
        raise FileNotFoundError(
            f"Selected run is not a multi-ligand result: {selected_run_dir}. Use the Multi-Ligand mode output."
        )

    safe_ligand = re.sub(r"[^A-Za-z0-9_.-]+", "_", selected_ligand)
    work_dir = temp_root / f"multi_ligand_{dtype}_{safe_ligand}_{selected_run_name}"
    out_path = _next_unique_png_path(output_dir, output_stem or f"{dtype}_{safe_ligand}_{selected_run_name}_{timestamp_token()}")

    try:
        run_pipeline(
            receptor_id=dtype,
            ligand_name=selected_ligand,
            run_dir=selected_run_dir,
            output_png=out_path,
            work_dir=work_dir,
            dpi=dpi,
            style_preset=otofigure_style,
            ray_trace=otofigure_ray_trace,
            options=otofigure_options,
            preview_mode=preview_mode,
            on_process_start=(process_hooks or {}).get("on_process_start"),
            on_process_end=(process_hooks or {}).get("on_process_end"),
        )
        return out_path, [selected_run_name]
    finally:
        shutil.rmtree(work_dir, ignore_errors=True)


def _get_render_panel_builder(render_mode: str):
    normalized = _normalize_render_mode(render_mode)
    if normalized == REPORT_RENDER_MODE_MULTI_LIGAND:
        return _render_dtype_multi_ligand_panel
    if normalized == REPORT_RENDER_MODE_OTOFIGURE:
        return _render_dtype_otofigure_panel
    return _render_dtype_panel


def _build_report_doc(
    report_root: Path,
    source_dir: Path,
    output_root: Path,
    *,
    images_root: Path | None = None,
    selected_images: list[Path] | None = None,
    figure_captions: dict[Path, str] | None = None,
    figure_start_number: int = 1,
    extra_sections: list[dict[str, str]] | None = None,
) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    _, render_dir, plot_dir, doc_path = _report_output_paths(output_root)

    selected = [path.resolve() for path in (selected_images or []) if path.exists() and path.is_file()]

    source_metadata = _load_source_metadata(source_dir)
    main_type = str(source_metadata.get("main_type") or "").strip()
    receptor_order = [str(item) for item in (source_metadata.get("receptor_order") or []) if str(item)]
    ligand_order = [str(item) for item in (source_metadata.get("ligand_order") or []) if str(item)]
    receptor_labels = {str(k): str(v) for k, v in (source_metadata.get("receptor_labels") or {}).items()}
    ligand_labels = {str(k): str(v) for k, v in (source_metadata.get("ligand_labels") or {}).items()}
    resolved_figure_start = _normalize_positive_int(
        figure_start_number or source_metadata.get("figure_start_number", 1),
        default=1,
        min_value=1,
        max_value=999,
    )
    resolved_extra_sections = _normalize_extra_sections(extra_sections or source_metadata.get("extra_sections", []))
    caption_overrides: dict[Path, str] = {}
    source_caption_map = _normalize_caption_map(source_metadata.get("figure_caption_overrides", {}))
    for raw_path, raw_caption in source_caption_map.items():
        candidate = Path(str(raw_path)).expanduser()
        if not candidate.is_absolute():
            ws = (WORKSPACE_DIR / candidate).resolve()
            if ws.exists():
                candidate = ws
            else:
                candidate = (BASE / candidate).resolve()
        else:
            candidate = candidate.resolve()
        caption_text = str(raw_caption or "").strip()
        if caption_text:
            caption_overrides[candidate] = caption_text
    for raw_path, raw_caption in (figure_captions or {}).items():
        key_path = Path(raw_path).resolve()
        caption_text = str(raw_caption or "").strip()
        if caption_text:
            caption_overrides[key_path] = caption_text

    if selected:
        ordered_images = selected
    elif images_root is not None:
        all_images = _list_images_under_root(images_root, category="report")
        collected: list[Path] = []
        for row in all_images:
            raw_path = str(row.get("path") or "").strip()
            if not raw_path:
                continue
            ws_candidate = (WORKSPACE_DIR / raw_path).resolve()
            if ws_candidate.exists() and ws_candidate.is_file():
                collected.append(ws_candidate)
            else:
                candidate = (BASE / raw_path).resolve()
                if candidate.exists() and candidate.is_file():
                    collected.append(candidate)
        ordered_images = collected
    else:
        ordered_images = sorted(render_dir.glob("*.png")) + sorted(plot_dir.glob("*.png"))

    if not ordered_images:
        raise HTTPException(status_code=400, detail="No images found for report.")

    selected_order_index = {path.resolve(): idx for idx, path in enumerate(ordered_images)}
    receptor_order_index = {name.upper(): idx for idx, name in enumerate(receptor_order)}

    def _image_category(path_obj: Path) -> str:
        parent_name = path_obj.parent.name.lower()
        text = str(path_obj).replace("\\", "/").lower()
        if parent_name == "render_images" or "/render_images/" in text:
            return "render"
        if parent_name == "plots" or "/plots/" in text:
            return "plot"
        return "other"

    def _plot_order_index(path_obj: Path) -> int:
        stem = path_obj.stem.lower()
        for idx, token in enumerate(REPORT_PLOT_ORDER_BY_NAME):
            if token in stem:
                return idx
        return len(REPORT_PLOT_ORDER_BY_NAME)

    def _render_order_key(path_obj: Path) -> tuple[int, int, str]:
        stem = path_obj.stem
        match = re.match(r"^(D\d+)_", stem, re.IGNORECASE)
        receptor_id = str(match.group(1) if match else "").upper()
        return (
            receptor_order_index.get(receptor_id, 10**6),
            selected_order_index.get(path_obj.resolve(), 10**6),
            stem.lower(),
        )

    render_images = [p for p in ordered_images if _image_category(p) == "render"]
    plot_images = [p for p in ordered_images if _image_category(p) == "plot"]
    other_images = [p for p in ordered_images if _image_category(p) not in {"render", "plot"}]
    render_images.sort(key=_render_order_key)
    plot_images.sort(
        key=lambda p: (
            _plot_order_index(p),
            selected_order_index.get(p.resolve(), 10**6),
            p.stem.lower(),
        )
    )
    other_images.sort(key=lambda p: selected_order_index.get(p.resolve(), 10**6))
    ordered_images = render_images + plot_images + other_images

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    usable_width = float((section.page_width - section.left_margin - section.right_margin) / 914400)
    image_width = max(5.9, min(6.8, usable_width))

    title_text = f"{main_type.upper()} DOCKING REPORT" if main_type else "DOCKING REPORT"

    heading = doc.add_heading(title_text, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Dock root: {to_display_path(report_root)}")
    doc.add_paragraph(f"Source root: {to_display_path(source_dir)}")
    doc.add_paragraph(f"Output root: {to_display_path(output_root)}")
    if images_root is not None:
        doc.add_paragraph(f"Images root: {to_display_path(images_root)}")
    doc.add_paragraph(f"Generated at: {time.strftime('%Y-%m-%d %H:%M:%S')}")

    def _default_caption(image_path: Path, figure_no: int) -> str:
        text = image_path.stem.lower()
        category = _image_category(image_path)
        receptor_match = re.match(r"^(D\d+)_", image_path.stem, re.IGNORECASE)
        receptor_id = str(receptor_match.group(1) if receptor_match else "").upper()
        receptor_label = receptor_labels.get(receptor_id, receptor_id or "receptor")
        ligands = [ligand_labels.get(name, _prettify_label(name, trim_run_suffix=True)) for name in ligand_order[:4]]
        ligand_clause = ", ".join(f"{chr(65 + idx)}) {name}" for idx, name in enumerate(ligands))
        receptor_group = ", ".join(receptor_labels.get(item, item) for item in receptor_order) or "selected receptors"
        if category == "render":
            if ligand_clause:
                return f"Figure {figure_no}. Final docking and interaction visualization for {receptor_label}. {ligand_clause}"
            return f"Figure {figure_no}. Final docking and interaction visualization for {receptor_label}."
        if "affinity_boxplot" in text:
            target = f"{main_type} receptor subtypes" if main_type else receptor_group
            ligand_names = ", ".join(ligands) if ligands else "selected ligands"
            return (
                f"Figure {figure_no}. Box plots illustrating docking-score distributions for {ligand_names} "
                f"against {target}. (Top) Mean binding affinities (kcal/mol) with standard deviations are tabulated."
            )
        if "run_frequency_heatmap" in text:
            target = f"{main_type} receptors" if main_type else "selected receptors"
            ligand_names = ", ".join(ligands) if ligands else "selected ligands"
            return (
                f"Figure {figure_no}. Interacting residues obtained via PLIP across {target} for {ligand_names}. "
                "Cell values (0-5) indicate in how many runs each residue was detected as interacting."
            )
        if "common_residue_heatmap" in text:
            target = f"{main_type} receptor subtypes" if main_type else "selected receptor subtypes"
            ligand_names = ", ".join(ligands) if ligands else "selected ligands"
            return (
                f"Figure {figure_no}. Common interacting residues (frequency = 5 for each run) by dominant interaction type "
                f"across runs for {ligand_names} in {target}."
            )
        if "interaction_stacked_bar" in text:
            target = f"{main_type} receptor subtypes" if main_type else "selected receptor subtypes"
            ligand_names = ", ".join(ligands) if ligands else "selected ligands"
            return (
                f"Figure {figure_no}. Common stacked counts of PLIP interaction instances for {ligand_names} "
                f"across {target}, using the best-scoring run for each receptor-ligand pair."
            )
        return f"Figure {figure_no}. {image_path.name}"

    figure_counter = [resolved_figure_start]

    def _add_centered_image(image_path: Path) -> None:
        resolved = image_path.resolve()
        caption_text = caption_overrides.get(resolved) or _default_caption(image_path, figure_counter[0])
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(image_path), width=Inches(image_width))
        caption = doc.add_paragraph(caption_text)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        figure_counter[0] += 1

    doc.add_heading(REPORT_TEMPLATE_HEADINGS[0], level=1)
    doc.add_heading(REPORT_TEMPLATE_HEADINGS[1], level=1)

    current_group = ""
    for image_path in ordered_images:
        category = _image_category(image_path)
        if category != current_group:
            if category == "render":
                doc.add_heading("Docking figures with Interacting residues", level=2)
            elif category == "plot":
                doc.add_heading("Predefined Plots", level=2)
            else:
                doc.add_heading("Additional Figures", level=2)
            current_group = category
        _add_centered_image(image_path)

    for section_row in resolved_extra_sections:
        title = str(section_row.get("title") or "").strip()
        body = str(section_row.get("body") or "").strip()
        if title:
            doc.add_heading(title, level=1)
        if body:
            doc.add_paragraph(body)

    doc.add_heading(REPORT_TEMPLATE_HEADINGS[2], level=1)
    doc.add_heading(REPORT_TEMPLATE_HEADINGS[3], level=1)

    doc_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(doc_path)
    return doc_path


@router.get("/api/reports/list")
def list_reports(root_path: str = "", source_path: str = "", output_path: str = "", linked_path: str = "") -> JSONResponse:
    report_root = _resolve_report_root(root_path)
    # Graceful fallback: if source_path is invalid, use the default source
    try:
        source_dir = _resolve_report_source(report_root, source_path)
    except HTTPException:
        source_dir = _default_report_source(report_root)
    # Same fallback for output path
    try:
        output_root = _resolve_report_output_root(report_root, source_dir, output_path)
    except HTTPException:
        output_root = (source_dir / "report_outputs").resolve()
        output_root.mkdir(parents=True, exist_ok=True)

    _ = linked_path
    receptor_rows = _collect_receptor_rows(source_dir)
    receptor_ids, ligand_names = _collect_entities_from_rows(receptor_rows)
    source_metadata = _load_source_metadata(source_dir, receptor_ids, ligand_names)
    receptors = _apply_source_metadata_to_rows(receptor_rows, source_metadata)
    source_folders = _list_source_folders(report_root, selected_source=source_dir)
    linked_error = "" if receptors else "No receptor/ligand/run detected in selected source."
    ready_rows = [row for row in receptors if row.get("ready")]

    payload = _collect_report_outputs(output_root)
    payload["root_path"] = to_display_path(report_root)
    payload["source_path"] = to_display_path(source_dir)
    payload["output_path"] = to_display_path(output_root)
    payload["source_folders"] = source_folders
    payload["linked_root_path"] = to_display_path(source_dir)
    payload["linked_roots"] = source_folders
    payload["linked_error"] = linked_error
    payload["source_metadata"] = source_metadata
    payload["receptors"] = receptors
    payload["default_receptors"] = [row["id"] for row in ready_rows]
    payload["dock_validation"] = {
        "inside_dock": True,
        "linked_roots_found": len([row for row in source_folders if row.get("receptor_count", 0) > 0]),
        "selected_receptors": len(receptors),
        "selected_ready_receptors": len(ready_rows),
        "selected_docking_count": sum(
            int(row.get("runs_per_ligand_min") or 0) for row in ready_rows
        ),
    }
    payload["images"] = payload["render_images"] + payload["plot_images"]
    return JSONResponse(payload)


@router.get("/api/reports/preview")
def get_report_preview(
    root_path: str = "",
    source_path: str = "",
    receptor_id: str = "",
    run_name: str = "",
    render_mode: str = "",
) -> JSONResponse:
    report_root = _resolve_report_root(root_path)
    source_dir = _resolve_report_source(report_root, source_path)
    try:
        payload = _resolve_preview_context(
            source_dir,
            receptor_id=receptor_id,
            run_name=run_name,
            render_mode=render_mode,
        )
    except FileNotFoundError as exc:
        return JSONResponse(
            {
                "available": False,
                "message": str(exc),
                "receptor_id": str(receptor_id or "").strip(),
                "render_mode": _normalize_render_mode(render_mode),
            }
        )
    payload["root_path"] = to_display_path(report_root)
    payload["source_path"] = to_display_path(source_dir)
    return JSONResponse(payload)


@router.get("/api/reports/images")
def list_report_images(
    root_path: str = "",
    source_path: str = "",
    output_path: str = "",
    images_root_path: str = "",
) -> JSONResponse:
    report_root = _resolve_report_root(root_path)
    source_dir = _resolve_report_source(report_root, source_path)
    output_root = _resolve_report_output_root(report_root, source_dir, output_path)
    images_root = _resolve_report_images_root(report_root, output_root, images_root_path)
    images = _list_images_under_root(images_root, category="report")
    return JSONResponse(
        {
            "root_path": to_display_path(report_root),
            "source_path": to_display_path(source_dir),
            "output_path": to_display_path(output_root),
            "images_root_path": to_display_path(images_root),
            "images": images,
            "total": len(images),
        }
    )


@router.get("/api/reports/root-metadata")
def get_report_root_metadata(root_path: str = "", source_path: str = "") -> JSONResponse:
    report_root = _resolve_report_root(root_path)
    source_dir = _resolve_report_source(report_root, source_path)
    receptors, ligands = _collect_source_entities(source_dir)
    metadata = _load_source_metadata(source_dir, receptors, ligands)
    return JSONResponse(
        {
            "root_path": to_display_path(report_root),
            "source_path": to_display_path(source_dir),
            "metadata": metadata,
        }
    )


@router.get("/api/reports/doc-config")
def get_report_doc_config(root_path: str = "", source_path: str = "") -> JSONResponse:
    report_root = _resolve_report_root(root_path)
    source_dir = _resolve_report_source(report_root, source_path)
    receptors, ligands = _collect_source_entities(source_dir)
    metadata = _load_source_metadata(source_dir, receptors, ligands)
    return JSONResponse(
        {
            "root_path": to_display_path(report_root),
            "source_path": to_display_path(source_dir),
            "figure_start_number": metadata.get("figure_start_number", 1),
            "extra_sections": metadata.get("extra_sections", []),
            "figure_caption_overrides": metadata.get("figure_caption_overrides", {}),
        }
    )


@router.post("/api/reports/doc-config")
def save_report_doc_config(payload: dict[str, Any]) -> JSONResponse:
    report_root = _resolve_report_root(str(payload.get("root_path") or ""))
    source_dir = _resolve_report_source(report_root, str(payload.get("source_path") or ""))
    receptors, ligands = _collect_source_entities(source_dir)
    metadata = _load_source_metadata(source_dir, receptors, ligands)

    figure_start_number = _normalize_positive_int(
        payload.get("figure_start_number", metadata.get("figure_start_number", 1)),
        default=1,
        min_value=1,
        max_value=999,
    )
    extra_sections = _normalize_extra_sections(payload.get("extra_sections", metadata.get("extra_sections", [])))
    figure_caption_overrides = _normalize_caption_map(
        payload.get("figure_caption_overrides", metadata.get("figure_caption_overrides", {}))
    )

    saved = _save_source_metadata(
        source_dir,
        main_type=str(metadata.get("main_type") or ""),
        receptor_labels={str(k): str(v) for k, v in (metadata.get("receptor_labels") or {}).items()},
        ligand_labels={str(k): str(v) for k, v in (metadata.get("ligand_labels") or {}).items()},
        receptor_order=[str(v) for v in (metadata.get("receptor_order") or []) if str(v)],
        ligand_order=[str(v) for v in (metadata.get("ligand_order") or []) if str(v)],
        figure_start_number=figure_start_number,
        extra_sections=extra_sections,
        figure_caption_overrides=figure_caption_overrides,
    )
    return JSONResponse(
        {
            "ok": True,
            "figure_start_number": saved.get("figure_start_number", 1),
            "extra_sections": saved.get("extra_sections", []),
            "figure_caption_overrides": saved.get("figure_caption_overrides", {}),
            "metadata": saved,
        }
    )


@router.post("/api/reports/root-metadata")
def save_report_root_metadata(payload: dict[str, Any]) -> JSONResponse:
    report_root = _resolve_report_root(str(payload.get("root_path") or ""))
    source_dir = _resolve_report_source(report_root, str(payload.get("source_path") or ""))
    receptors, ligands = _collect_source_entities(source_dir)
    existing_metadata = _load_source_metadata(source_dir, receptors, ligands)

    if bool(payload.get("reset")):
        meta_file = _metadata_file_for_source(source_dir)
        if meta_file.exists() and meta_file.is_file():
            meta_file.unlink()
        metadata = _load_source_metadata(source_dir, receptors, ligands)
        return JSONResponse({"ok": True, "metadata": metadata})

    main_type = str(payload.get("main_type") or "").strip()
    receptor_map_raw = _normalize_label_map(payload.get("receptor_labels"))
    ligand_map_raw = _normalize_label_map(payload.get("ligand_labels"))
    receptor_order = _normalize_order_list(payload.get("receptor_order"), receptors)
    ligand_order = _normalize_order_list(payload.get("ligand_order"), ligands)

    receptor_labels = {
        receptor_id: receptor_map_raw.get(receptor_id, _prettify_label(receptor_id))
        for receptor_id in receptors
    }
    ligand_labels = {
        ligand_name: ligand_map_raw.get(ligand_name, _prettify_label(ligand_name, trim_run_suffix=True))
        for ligand_name in ligands
    }
    figure_start_number = _normalize_positive_int(
        payload.get("figure_start_number", existing_metadata.get("figure_start_number", 1)),
        default=1,
        min_value=1,
        max_value=999,
    )
    extra_sections = _normalize_extra_sections(payload.get("extra_sections", existing_metadata.get("extra_sections", [])))
    figure_caption_overrides = _normalize_caption_map(
        payload.get("figure_caption_overrides", existing_metadata.get("figure_caption_overrides", {}))
    )
    metadata = _save_source_metadata(
        source_dir,
        main_type=main_type,
        receptor_labels=receptor_labels,
        ligand_labels=ligand_labels,
        receptor_order=receptor_order,
        ligand_order=ligand_order,
        figure_start_number=figure_start_number,
        extra_sections=extra_sections,
        figure_caption_overrides=figure_caption_overrides,
    )
    return JSONResponse({"ok": True, "metadata": metadata})


@router.post("/api/reports/source/delete")
def delete_report_source(payload: dict[str, Any]) -> JSONResponse:
    report_root = _resolve_report_root(str(payload.get("root_path") or ""))
    source_dir = _resolve_report_source(report_root, str(payload.get("source_path") or ""))
    report_root = report_root.resolve()
    source_dir = source_dir.resolve()

    if source_dir == report_root:
        raise HTTPException(status_code=400, detail="Cannot delete report root itself.")
    if source_dir.parent != report_root:
        raise HTTPException(status_code=400, detail="Only first-level source folders can be deleted.")
    if not source_dir.exists() or not source_dir.is_dir():
        raise HTTPException(status_code=404, detail="Source folder not found.")

    deleted_path = to_display_path(source_dir)
    shutil.rmtree(source_dir, ignore_errors=False)

    remaining = sorted(
        [p for p in report_root.iterdir() if p.is_dir() and not p.name.startswith(".") and p.name not in {"report_outputs", "__pycache__"}],
        key=lambda p: p.name.lower(),
    )
    next_source = to_display_path(remaining[0]) if remaining else to_display_path(report_root)
    return JSONResponse(
        {
            "ok": True,
            "deleted": deleted_path,
            "next_source_path": next_source,
        }
    )


@router.post("/api/reports/images/delete-all")
def delete_all_report_images(payload: dict[str, Any]) -> JSONResponse:
    report_root = _resolve_report_root(str(payload.get("root_path") or ""))
    source_dir = _resolve_report_source(report_root, str(payload.get("source_path") or ""))
    output_root = _resolve_report_output_root(report_root, source_dir, str(payload.get("output_path") or ""))
    _, render_dir, plot_dir, _ = _report_output_paths(output_root)

    scope = str(payload.get("scope") or "all").strip().lower()
    target_dirs: list[Path] = []
    if scope in {"render", "render_images"}:
        target_dirs = [render_dir]
        scope_label = "render"
    elif scope in {"plot", "plots", "graph", "graphs"}:
        target_dirs = [plot_dir]
        scope_label = "plot"
    elif scope in {"all", ""}:
        target_dirs = [render_dir, plot_dir]
        scope_label = "all"
    else:
        raise HTTPException(status_code=400, detail="Invalid delete scope.")

    deleted: list[str] = []
    for directory in target_dirs:
        if not directory.exists() or not directory.is_dir():
            continue
        for image_path in directory.iterdir():
            if not image_path.is_file() or not _is_report_image_file(image_path):
                continue
            _delete_image_artifacts(image_path)
            rel_path = relative_to_base(image_path)
            if rel_path:
                deleted.append(rel_path)

    return JSONResponse(
        {
            "ok": True,
            "scope": scope_label,
            "deleted_count": len(deleted),
            "deleted": deleted[:200],
        }
    )


@router.post("/api/reports/image/delete")
def delete_report_image(payload: dict[str, Any]) -> JSONResponse:
    report_root = _resolve_report_root(str(payload.get("root_path") or ""))
    source_dir = _resolve_report_source(report_root, str(payload.get("source_path") or ""))
    output_root = _resolve_report_output_root(report_root, source_dir, str(payload.get("output_path") or ""))
    images_root = _resolve_report_images_root(report_root, output_root, str(payload.get("images_root_path") or ""))
    target = _resolve_report_image_path(report_root, images_root, str(payload.get("path") or ""))
    if not target.exists():
        raise HTTPException(status_code=404, detail="Image not found")
    if not target.is_file():
        raise HTTPException(status_code=400, detail="Invalid image path")
    _delete_image_artifacts(target)
    return JSONResponse({"ok": True, "deleted": relative_to_base(target)})


@router.get("/api/reports/image/{path:path}")
def serve_report_image(path: str):
    requested = Path(str(path or "").strip()).expanduser()
    if not requested.is_absolute():
        ws_candidate = (WORKSPACE_DIR / requested).resolve()
        if ws_candidate.exists():
            target_path = ws_candidate
        else:
            target_path = (BASE / requested).resolve()
    else:
        target_path = requested.resolve()

    if target_path != DOCK_DIR_RESOLVED and DOCK_DIR_RESOLVED not in target_path.parents:
        raise HTTPException(status_code=404, detail="Image not found")
    if target_path.suffix.lower() not in REPORT_IMAGE_EXTENSIONS:
        raise HTTPException(status_code=404, detail="Image not found")
    if not target_path.exists() or not target_path.is_file():
        raise HTTPException(status_code=404, detail="Image not found")
    return FileResponse(str(target_path))


@router.get("/api/reports/doc")
def serve_report_doc(root_path: str = "", source_path: str = "", output_path: str = ""):
    report_root = _resolve_report_root(root_path)
    source_dir = _resolve_report_source(report_root, source_path)
    output_root = _resolve_report_output_root(report_root, source_dir, output_path)
    _, _, _, doc_path = _report_output_paths(output_root)
    if not doc_path.exists():
        raise HTTPException(status_code=404, detail="Report document not found")
    return FileResponse(
        str(doc_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=doc_path.name,
    )


@router.post("/api/reports/graphs")
def trigger_graphs(payload: GraphPayload, background_tasks: BackgroundTasks) -> JSONResponse:
    global REPORT_STATE
    if REPORT_STATE.get("status") == "running":
        return JSONResponse({"error": "Another report task is already running."}, status_code=409)

    report_root = _resolve_report_root(payload.root_path)
    source_dir = _resolve_report_source(report_root, payload.source_path)
    output_root = _resolve_report_output_root(report_root, source_dir, payload.output_path)
    _ = payload.linked_path
    _, _, plot_dir, _ = _report_output_paths(output_root)
    plot_dir.mkdir(parents=True, exist_ok=True)
    receptor_rows = _collect_receptor_rows(source_dir)
    if not any(bool(row.get("ready")) for row in receptor_rows):
        fallback_source = _find_ready_report_source(report_root, source_dir)
        if fallback_source != source_dir:
            source_dir = fallback_source
            output_root = (source_dir / "report_outputs").resolve()
            output_root.mkdir(parents=True, exist_ok=True)
            _, _, plot_dir, _ = _report_output_paths(output_root)
            plot_dir.mkdir(parents=True, exist_ok=True)
            receptor_rows = _collect_receptor_rows(source_dir)
        if not any(bool(row.get("ready")) for row in receptor_rows):
            return JSONResponse({"error": "No receptor/ligand/run detected in selected source."}, status_code=400)

    requested = [script_id for script_id in payload.scripts if script_id in REPORT_PREDEFINED_PLOTS]
    selected = requested or list(REPORT_PREDEFINED_PLOTS.keys())

    REPORT_STATE["status"] = "running"
    REPORT_STATE["task"] = "plots"
    REPORT_STATE["progress"] = 0
    REPORT_STATE["total"] = len(selected)
    REPORT_STATE["message"] = "Generating predefined plots..."
    REPORT_STATE["errors"] = []
    REPORT_STATE["last_logs"] = []

    def run_plot_job(state: dict[str, Any], script_ids: list[str], root: Path, out_dir: Path) -> None:
        errors: list[str] = []
        logs: list[str] = []
        started_stamp = timestamp_token()

        for idx, script_id in enumerate(script_ids, start=1):
            if state.get("status") != "running":
                break
            spec = REPORT_PREDEFINED_PLOTS[script_id]
            state["message"] = f"[{idx}/{len(script_ids)}] {spec['label']}"
            tmp_out = out_dir / f".tmp_{script_id}_{time.time_ns()}"
            tmp_out.mkdir(parents=True, exist_ok=True)
            cmd = [
                sys.executable,
                "-m",
                spec["module"],
                "--root",
                str(root),
                "--out",
                str(tmp_out),
            ]
            proc = subprocess.run(
                cmd,
                cwd=str(BASE),
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
            )
            tail_line = (proc.stdout or "").strip().splitlines()[-1:] or [""]
            expected_file = tmp_out / spec["filename"]
            if proc.returncode != 0 or not expected_file.exists():
                detail = tail_line[0].strip() if tail_line else ""
                errors.append(f"{script_id} failed{(': ' + detail) if detail else ''}")
                logs.append(f"{script_id}: {tail_line[0]}")
            else:
                stem = Path(spec["filename"]).stem
                final_stem = f"{stem}_{started_stamp}_{idx:02d}"
                final_path = _next_unique_png_path(out_dir, final_stem)
                shutil.move(str(expected_file), str(final_path))
                logs.append(f"{script_id}: {final_path.name}")
            shutil.rmtree(tmp_out, ignore_errors=True)
            state["progress"] = idx

        state["errors"] = errors
        state["last_logs"] = logs[-10:]
        state["status"] = "idle"
        state["message"] = "Plots generated." if not errors else "Plots completed with errors."

    background_tasks.add_task(run_plot_job, REPORT_STATE, selected, source_dir, plot_dir)
    return JSONResponse(
        {
            "status": "started",
            "source_path": to_display_path(source_dir),
            "output_path": to_display_path(output_root),
        }
    )


@router.post("/api/reports/render")
def trigger_render(payload: RenderPayload, background_tasks: BackgroundTasks) -> JSONResponse:
    global REPORT_STATE
    if REPORT_STATE.get("status") == "running":
        return JSONResponse({"error": "Another report task is already running."}, status_code=409)
    try:
        render_mode = _normalize_render_mode(payload.render_mode)
    except ValueError as exc:
        return JSONResponse({"error": str(exc)}, status_code=400)

    report_root = _resolve_report_root(payload.root_path)
    source_dir = _resolve_report_source(report_root, payload.source_path)
    output_root = _resolve_report_output_root(report_root, source_dir, payload.output_path)
    _ = payload.linked_path
    _, render_dir, _, _ = _report_output_paths(output_root)
    render_dir.mkdir(parents=True, exist_ok=True)

    receptor_rows = _collect_receptor_rows(source_dir)
    receptor_ids, ligand_names = _collect_entities_from_rows(receptor_rows)
    source_metadata = _load_source_metadata(source_dir, receptor_ids, ligand_names)
    ligand_order_index = {
        str(name): idx
        for idx, name in enumerate(source_metadata.get("ligand_order") or [])
        if str(name)
    }
    receptor_inventory = _collect_receptor_inventory(source_dir)
    ready_receptors = [row["id"] for row in receptor_rows if row.get("ready")]
    receptor_index = {row["id"]: row for row in receptor_rows}
    selected = [item for item in payload.receptors if item in ready_receptors]
    if not selected:
        selected = ready_receptors
    if payload.is_preview and selected:
        selected = selected[:1]
    if not selected:
        fallback_source = _find_ready_report_source(report_root, source_dir)
        if fallback_source != source_dir:
            source_dir = fallback_source
            output_root = (source_dir / "report_outputs").resolve()
            output_root.mkdir(parents=True, exist_ok=True)
            _, render_dir, _, _ = _report_output_paths(output_root)
            render_dir.mkdir(parents=True, exist_ok=True)
            receptor_rows = _collect_receptor_rows(source_dir)
            receptor_ids, ligand_names = _collect_entities_from_rows(receptor_rows)
            source_metadata = _load_source_metadata(source_dir, receptor_ids, ligand_names)
            ligand_order_index = {
                str(name): idx
                for idx, name in enumerate(source_metadata.get("ligand_order") or [])
                if str(name)
            }
            receptor_inventory = _collect_receptor_inventory(source_dir)
            ready_receptors = [row["id"] for row in receptor_rows if row.get("ready")]
            receptor_index = {row["id"]: row for row in receptor_rows}
            selected = [item for item in payload.receptors if item in ready_receptors]
            if not selected:
                selected = ready_receptors
            if payload.is_preview and selected:
                selected = selected[:1]
        if not selected:
            return JSONResponse({"error": "No render-ready receptors found for selected source."}, status_code=400)

    is_preview_mode = bool(payload.is_preview)
    dpi = int(payload.dpi or 120)
    dpi = max(30, min(600, dpi))
    otofigure_style = str(payload.otofigure_style or "balanced").strip().lower() or "balanced"
    if otofigure_style not in {"balanced", "ligand_focus", "surface_focus"}:
        otofigure_style = "balanced"
    otofigure_render_engine = str(payload.otofigure_render_engine or "").strip().lower() or "ray"
    if otofigure_render_engine not in {"ray", "opengl", "fast_draw"}:
        otofigure_render_engine = "ray"
    otofigure_ray_trace = bool(payload.otofigure_ray_trace)
    if otofigure_render_engine == "ray":
        otofigure_ray_trace = True
    else:
        otofigure_ray_trace = False
    otofigure_background = str(payload.otofigure_background or "transparent").strip().lower() or "transparent"
    if otofigure_background not in {"transparent", "white"}:
        otofigure_background = "transparent"
    otofigure_protein_color = str(payload.otofigure_protein_color or "bluewhite").strip() or "bluewhite"
    otofigure_options = {
        "render_engine": otofigure_render_engine,
        "background": otofigure_background,
        "surface_enabled": bool(payload.otofigure_surface_enabled),
        "surface_opacity": max(0.0, min(1.0, float(payload.otofigure_surface_opacity))),
        "protein_color": otofigure_protein_color,
        "ligand_thickness": max(0.05, min(0.8, float(payload.otofigure_ligand_thickness))),
        "far_ratio": max(1, min(9, int(payload.otofigure_far_ratio))),
        "close_ratio": max(1, min(9, int(payload.otofigure_close_ratio))),
        "interaction_ratio": max(1, min(9, int(payload.otofigure_interaction_ratio))),
        "far_padding": max(0.0, min(0.5, float(payload.otofigure_far_padding))),
        "far_frame_margin": max(0.0, min(0.15, float(payload.otofigure_far_frame_margin))),
        "close_padding": max(0.0, min(1.0, float(payload.otofigure_close_padding))),
    }
    render_panel_builder = _get_render_panel_builder(render_mode)
    if render_mode == REPORT_RENDER_MODE_OTOFIGURE:
        render_mode_label = "OtoFigure"
    elif render_mode == REPORT_RENDER_MODE_MULTI_LIGAND:
        render_mode_label = "Multi-Ligand Panel"
    else:
        render_mode_label = "Classic"

    render_jobs: list[dict[str, str]] = []
    for dtype in selected:
        preferred_run = str(payload.run_by_receptor.get(dtype, "")).strip()
        preferred_ligand = str(payload.ligand_by_receptor.get(dtype, "")).strip()
        row = receptor_index.get(dtype) or {}
        run_options = {run_name for run_name in row.get("run_options", []) if run_name}
        if preferred_run and run_options and preferred_run not in run_options:
            preferred_run = str(row.get("default_run") or "")
        if not preferred_run:
            preferred_run = str(row.get("default_run") or "run1")
        render_jobs.append({"dtype": dtype, "run": preferred_run, "ligand": preferred_ligand})

    REPORT_STATE["status"] = "running"
    REPORT_STATE["task"] = "render"
    REPORT_STATE["progress"] = 0
    REPORT_STATE["total"] = len(render_jobs)
    REPORT_STATE["message"] = f"Generating {render_mode_label} render panels..."
    REPORT_STATE["errors"] = []
    REPORT_STATE["last_logs"] = []
    REPORT_STATE["cancel_requested"] = False
    REPORT_STATE["current_receptor"] = ""
    REPORT_STATE["current_ligand"] = ""
    REPORT_STATE["current_run"] = ""
    REPORT_STATE["render_mode"] = render_mode
    REPORT_STATE["active_subprocess_pid"] = None
    REPORT_STATE["active_subprocess_label"] = ""

    def run_render_job(
        state: dict[str, Any],
        receptor_data: dict[str, dict[str, list[tuple[str, Path]]]],
        out_dir: Path,
        jobs: list[dict[str, str]],
        render_dpi: int,
        preview_mode: bool,
        ligand_order: dict[str, int],
        render_builder,
        render_mode_name: str,
        render_mode_key: str,
        otofigure_style_name: str,
        otofigure_ray_enabled: bool,
        otofigure_options_data: dict[str, Any],
    ) -> None:
        errors: list[str] = []
        logs: list[str] = []
        temp_root = out_dir.parent / ".tmp_render"
        if temp_root.exists():
            shutil.rmtree(temp_root, ignore_errors=True)
        temp_root.mkdir(parents=True, exist_ok=True)
        started_stamp = timestamp_token()

        try:
            for idx, job in enumerate(jobs, start=1):
                if state.get("cancel_requested"):
                    break
                dtype = job.get("dtype", "")
                preferred_run = (job.get("run") or "").strip()
                preferred_ligand = (job.get("ligand") or "").strip()
                run_label = "all runs" if render_mode_key == REPORT_RENDER_MODE_OTOFIGURE else (preferred_run if preferred_run else "auto")
                state["current_receptor"] = dtype
                state["current_ligand"] = preferred_ligand
                state["current_run"] = preferred_run
                state["message"] = (
                    f"[{idx}/{len(jobs)}] Rendering {dtype} with {render_mode_name} mode (run: {run_label})..."
                )

                def on_process_start(proc: subprocess.Popen[str]) -> None:
                    state["active_subprocess_pid"] = int(proc.pid or 0) or None
                    state["active_subprocess_label"] = dtype

                def on_process_end(_proc: subprocess.Popen[str]) -> None:
                    state["active_subprocess_pid"] = None
                    state["active_subprocess_label"] = ""

                try:
                    render_started_at = time.perf_counter()
                    if render_mode_key == REPORT_RENDER_MODE_OTOFIGURE:
                        panel_stem = f"{dtype}_{render_mode}_multirun_{started_stamp}_{idx:02d}"
                    else:
                        panel_stem = f"{dtype}_{render_mode}_{preferred_run or 'run1'}_{started_stamp}_{idx:02d}"
                    builder_kwargs: dict[str, Any] = {
                        "output_stem": panel_stem,
                        "preview_mode": preview_mode,
                        "ligand_order_index": ligand_order,
                        "process_hooks": {
                            "on_process_start": on_process_start,
                            "on_process_end": on_process_end,
                        },
                    }
                    if render_mode_key == REPORT_RENDER_MODE_OTOFIGURE:
                        builder_kwargs.update(
                            {
                                "otofigure_style": otofigure_style_name,
                                "otofigure_ray_trace": otofigure_ray_enabled,
                                "otofigure_options": otofigure_options_data,
                            }
                        )
                    signature = inspect.signature(render_builder)
                    accepts_var_kwargs = any(
                        parameter.kind == inspect.Parameter.VAR_KEYWORD
                        for parameter in signature.parameters.values()
                    )
                    if not accepts_var_kwargs:
                        builder_kwargs = {
                            key: value
                            for key, value in builder_kwargs.items()
                            if key in signature.parameters
                        }
                    out_path, used_runs = render_builder(
                        dtype,
                        receptor_data,
                        out_dir,
                        temp_root,
                        render_dpi,
                        preferred_run,
                        preferred_ligand,
                        **builder_kwargs,
                    )
                    elapsed_seconds = max(0.001, time.perf_counter() - render_started_at)
                    _write_image_metadata(
                        out_path,
                        {
                            "version": 1,
                            "kind": "render",
                            "dtype": dtype,
                            "render_mode": render_mode_key,
                            "render_dpi": int(render_dpi),
                            "preview_mode": bool(preview_mode),
                            "requested_run": preferred_run,
                            "requested_ligand": preferred_ligand,
                            "used_runs": list(used_runs or []),
                            "elapsed_seconds": round(elapsed_seconds, 3),
                            "otofigure_style": otofigure_style_name if render_mode_key == REPORT_RENDER_MODE_OTOFIGURE else "",
                            "otofigure_ray_trace": bool(otofigure_ray_enabled) if render_mode_key == REPORT_RENDER_MODE_OTOFIGURE else None,
                            "otofigure_render_engine": str(otofigure_options_data.get("render_engine") or "") if render_mode_key == REPORT_RENDER_MODE_OTOFIGURE else "",
                            "otofigure_background": str(otofigure_options_data.get("background") or "") if render_mode_key == REPORT_RENDER_MODE_OTOFIGURE else "",
                        },
                    )
                    used_label = ", ".join(used_runs) if used_runs else "auto"
                    logs.append(f"{dtype}: {out_path.name} (used: {used_label}; {elapsed_seconds:.1f}s)")
                except Exception as exc:
                    if state.get("cancel_requested"):
                        logs.append(f"{dtype}: cancelled")
                        break
                    errors.append(f"{dtype}: {exc}")
                state["progress"] = idx
        finally:
            shutil.rmtree(temp_root, ignore_errors=True)
            state["errors"] = errors
            state["last_logs"] = logs[-10:]
            state["active_subprocess_pid"] = None
            state["active_subprocess_label"] = ""
            state["current_receptor"] = ""
            state["current_ligand"] = ""
            state["current_run"] = ""
            state["status"] = "idle"
            if state.get("cancel_requested"):
                state["message"] = "Render stopped."
            else:
                state["message"] = "Render completed." if not errors else "Render completed with errors."
            state["cancel_requested"] = False

    background_tasks.add_task(
        run_render_job,
        REPORT_STATE,
        receptor_inventory,
        render_dir,
        render_jobs,
        dpi,
        is_preview_mode,
        ligand_order_index,
        render_panel_builder,
        render_mode_label,
        render_mode,
        otofigure_style,
        otofigure_ray_trace,
        otofigure_options,
    )
    dpi_scale = max(0.5, float(dpi) / 120.0)
    if render_mode == REPORT_RENDER_MODE_OTOFIGURE:
        base_seconds = 20 if is_preview_mode else 55
        engine_multiplier = {
            "ray": 1.0,
            "opengl": 0.55,
            "fast_draw": 0.30,
        }.get(str(otofigure_options.get("render_engine") or "ray"), 1.0)
        base_seconds *= engine_multiplier
    else:
        base_seconds = 14 if is_preview_mode else 40
    expected_seconds = int(max(12, len(render_jobs) * base_seconds * dpi_scale))
    REPORT_STATE["expected_time"] = expected_seconds
    return JSONResponse({"status": "started", "expected_time": expected_seconds})


@router.post("/api/reports/render/stop")
def stop_render() -> JSONResponse:
    if REPORT_STATE.get("task") != "render" or REPORT_STATE.get("status") not in {"running", "stopping"}:
        return JSONResponse({"status": "idle", "message": "No active render task."})

    REPORT_STATE["cancel_requested"] = True
    REPORT_STATE["status"] = "stopping"
    REPORT_STATE["message"] = "Stopping render..."

    active_pid = REPORT_STATE.get("active_subprocess_pid")
    if active_pid:
        try:
            os.killpg(int(active_pid), signal.SIGTERM)
        except ProcessLookupError:
            pass
        except OSError:
            pass

    return JSONResponse(
        {
            "status": REPORT_STATE.get("status", "stopping"),
            "message": REPORT_STATE.get("message", "Stopping render..."),
        }
    )


@router.post("/api/reports/compile")
def compile_report(payload: ReportCompilePayload) -> JSONResponse:
    global REPORT_STATE
    if REPORT_STATE.get("status") == "running":
        return JSONResponse({"error": "Another report task is already running."}, status_code=409)

    report_root = _resolve_report_root(payload.root_path)
    source_dir = _resolve_report_source(report_root, payload.source_path)
    output_root = _resolve_report_output_root(report_root, source_dir, payload.output_path)
    images_root = _resolve_report_images_root(report_root, output_root, payload.images_root_path)
    receptors, ligands = _collect_source_entities(source_dir)
    source_metadata = _load_source_metadata(source_dir, receptors, ligands)

    selected_images: list[Path] = []
    selected_figure_captions: dict[Path, str] = {}
    seen_selected: set[Path] = set()
    caption_map_raw: dict[str, str] = {}
    caption_map_raw.update(_normalize_caption_map(source_metadata.get("figure_caption_overrides", {})))
    caption_map_raw.update(_normalize_caption_map(payload.figure_captions or {}))
    for raw_image in payload.selected_images:
        candidate = _resolve_report_image_path(report_root, images_root, raw_image)
        if not candidate.exists() or not candidate.is_file():
            raise HTTPException(status_code=400, detail=f"Selected image not found: {raw_image}")
        if candidate in seen_selected:
            continue
        seen_selected.add(candidate)
        selected_images.append(candidate)
        relative_key = relative_to_base(candidate)
        caption_text = str(
            caption_map_raw.get(str(raw_image), "")
            or caption_map_raw.get(relative_key, "")
            or caption_map_raw.get(to_display_path(candidate), "")
        ).strip()
        if caption_text:
            selected_figure_captions[candidate.resolve()] = caption_text

    figure_start_number = _normalize_positive_int(
        payload.figure_start_number or source_metadata.get("figure_start_number", 1),
        default=1,
        min_value=1,
        max_value=999,
    )
    extra_sections = _normalize_extra_sections(payload.extra_sections or source_metadata.get("extra_sections", []))

    saved_meta = _save_source_metadata(
        source_dir,
        main_type=str(source_metadata.get("main_type") or ""),
        receptor_labels={str(k): str(v) for k, v in (source_metadata.get("receptor_labels") or {}).items()},
        ligand_labels={str(k): str(v) for k, v in (source_metadata.get("ligand_labels") or {}).items()},
        receptor_order=[str(v) for v in (source_metadata.get("receptor_order") or []) if str(v)],
        ligand_order=[str(v) for v in (source_metadata.get("ligand_order") or []) if str(v)],
        figure_start_number=figure_start_number,
        extra_sections=extra_sections,
        figure_caption_overrides=caption_map_raw,
    )
    source_metadata = saved_meta

    REPORT_STATE["status"] = "running"
    REPORT_STATE["task"] = "compile"
    REPORT_STATE["progress"] = 0
    REPORT_STATE["total"] = 1
    REPORT_STATE["message"] = "Compiling report document..."
    REPORT_STATE["errors"] = []
    REPORT_STATE["last_logs"] = []

    try:
        doc_path = _build_report_doc(
            report_root,
            source_dir,
            output_root,
            images_root=images_root,
            selected_images=selected_images,
            figure_captions=selected_figure_captions,
            figure_start_number=figure_start_number,
            extra_sections=extra_sections,
        )
        rel_doc = relative_to_base(doc_path)
        REPORT_STATE["progress"] = 1
        REPORT_STATE["message"] = "Report document created."
        REPORT_STATE["last_logs"] = [f"doc: {doc_path.name}"]
        return JSONResponse(
            {
                "status": "completed",
                "doc_path": rel_doc,
                "images_root_path": to_display_path(images_root),
                "selected_count": len(selected_images),
                "download_url": (
                    f"/api/reports/doc?root_path={to_display_path(report_root)}"
                    f"&source_path={to_display_path(source_dir)}"
                    f"&output_path={to_display_path(output_root)}"
                ),
            }
        )
    except HTTPException as exc:
        REPORT_STATE["errors"] = [str(exc.detail)]
        REPORT_STATE["message"] = "Report compilation failed."
        raise
    except Exception as exc:
        REPORT_STATE["errors"] = [str(exc)]
        REPORT_STATE["message"] = "Report compilation failed."
        raise HTTPException(status_code=500, detail=f"Report compilation failed: {exc}")
    finally:
        REPORT_STATE["status"] = "idle"


@router.get("/api/reports/status")
def get_report_status() -> JSONResponse:
    state = {
        "status": REPORT_STATE.get("status", "idle"),
        "task": REPORT_STATE.get("task", ""),
        "progress": REPORT_STATE.get("progress", 0),
        "total": REPORT_STATE.get("total", 0),
        "expected_time": REPORT_STATE.get("expected_time", 0),
        "message": REPORT_STATE.get("message", ""),
        "errors": REPORT_STATE.get("errors", []),
        "last_logs": REPORT_STATE.get("last_logs", []),
        "cancel_requested": bool(REPORT_STATE.get("cancel_requested")),
        "current_receptor": REPORT_STATE.get("current_receptor", ""),
        "current_ligand": REPORT_STATE.get("current_ligand", ""),
        "current_run": REPORT_STATE.get("current_run", ""),
        "render_mode": REPORT_STATE.get("render_mode", ""),
        "active_subprocess_pid": REPORT_STATE.get("active_subprocess_pid"),
        "active_subprocess_label": REPORT_STATE.get("active_subprocess_label", ""),
    }
    return JSONResponse(state)
