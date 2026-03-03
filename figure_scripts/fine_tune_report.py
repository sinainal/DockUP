"""
Fine-Tunes the Dimer Report.
Updates specific numerical values and Table 2 statistics with verified Dimer data.
"""
from pathlib import Path
from docx import Document

SOURCE_DOC = Path("dimer_report.docx")
OUT_DOC = Path("dimer_report_fined.docx")

# Verified Dimer Stats (Mean +/- SD)
STATS = {
    "D1": {"PET": "-7.86 ± 0.05", "PS": "-6.60 ± 0.22", "PP": "-3.98 ± 0.04", "PE": "-3.10 ± 0.00"},
    "D2": {"PET": "-9.44 ± 0.05", "PS": "-7.68 ± 0.04", "PP": "-4.70 ± 0.00", "PE": "-3.50 ± 0.00"},
    "D3": {"PET": "-8.30 ± 0.00", "PS": "-7.90 ± 0.00", "PP": "-4.40 ± 0.00", "PE": "-3.20 ± 0.00"},
    "D4": {"PET": "-9.44 ± 0.05", "PS": "-6.64 ± 0.22", "PP": "-4.32 ± 0.04", "PE": "-3.20 ± 0.00"},
    "D5": {"PET": "-8.32 ± 0.13", "PS": "-8.56 ± 0.05", "PP": "-4.60 ± 0.00", "PE": "-3.50 ± 0.00"},
}

def update_table(doc):
    """Updates Table 2 with new stats."""
    found = False
    for table in doc.tables:
        # Check headers to identify Table 2
        # Usually Row 0: Receptor, PET, PS, PP, PE
        try:
            headers = [c.text.strip() for c in table.rows[0].cells]
            if "Receptor" in headers and "PET" in headers:
                print("Found Data Table. Updating...")
                
                # Column indices
                try:
                    col_pet = headers.index("PET")
                    col_ps = headers.index("PS")
                    col_pp = headers.index("PP")
                    col_pe = headers.index("PE")
                except ValueError:
                    # Maybe headers use "PET dimer" if replaced?
                    # Since we did Strict Mode (Monomer->Dimer), header "PET" -> "PET" (preserved).
                    # Actually, we did Strict Mode on `dimer_report.docx`.
                    # Wait, strict mode replaced "Monomer"->"Dimer".
                    # Did it replace "PET" -> "PET dimer"? NO.
                    # So headers should be "PET", "PS"...
                    col_pet = headers.index("PET")
                    col_ps = headers.index("PS")
                    col_pp = headers.index("PP")
                    col_pe = headers.index("PE")

                for row in table.rows[1:]:
                    cells = row.cells
                    rec = cells[0].text.strip() # D1, D2..
                    
                    if rec in STATS:
                        # Update cells (preserving formatting if possible, else text replace)
                        cells[col_pet].text = STATS[rec]["PET"]
                        cells[col_ps].text = STATS[rec]["PS"]
                        cells[col_pp].text = STATS[rec]["PP"]
                        cells[col_pe].text = STATS[rec]["PE"]
                found = True
                break
        except IndexError:
            continue
            
    if not found:
        print("Warning: Data Table NOT found.")

def update_text(doc):
    """Updates specific text values."""
    # Target: "PET exhibits its highest binding affinity towards the D2 receptor (-7.2 kcal/mol)"
    # Dimer D2 PET is -9.44.
    
    # Target 2: Generic ranges if any.
    
    for p in doc.paragraphs:
        if "PET exhibits its highest binding affinity towards the D2 receptor" in p.text:
            print(f"Found Target Paragraph: {p.text[:50]}...")
            # We want to replace the value in parentheses roughly
            # Regex or direct replace
            if "-7.2" in p.text:
                p.text = p.text.replace("-7.2", "-9.44")
                print("  Updated -7.2 -> -9.44")
            # Might be other values (e.g. D3 is -8.30)
            
        # Check for other mentions
    # Target 3: Interaction Residues
    # Legacy text to find: "residues corresponding to this framework are distinctly repeated"
    # Parentheses to replace: (PHE... TRP309R (D5))
    NEW_INTERACTIONS = "PHE288F–ILE104F–LEU190F–VAL317F–TRP285F (D1); TRP386A–PHE390A–PHE198A–PHE389A–PHE382A (D2); PHE345A–PHE346A–HIS349A–ILE183A–VAL189A (D3); VAL87A–LEU111A–PHE91A–PHE410A–LEU187A (D4); PHE312R–VAL345R–PHE313R–TRP309R–ILE121R (D5)"
    
    for p in doc.paragraphs:
        if "residues corresponding to this framework are distinctly repeated" in p.text:
            print("Found Interaction Paragraph.")
            # Crude find/replace of the parenthetical block might be risky if we don't know exact text.
            # But the structure is "repeated (" ... "). This recurring"
            # Let's find the start "(" and end ")."
            try:
                start_idx = p.text.find("repeated (")
                end_idx = p.text.find("This recurring")
                if start_idx != -1 and end_idx != -1:
                    # Look for the closing paren before "This recurring"
                    # The legacy text ends with "... (D5)). This"
                    # We can construct the new text.
                    prefix = p.text[:start_idx + len("repeated ")]
                    suffix = p.text[end_idx:]
                    # Check what we are replacing
                    old_segment = p.text[start_idx:end_idx]
                    print(f"Replacing segment: {old_segment[:30]}...{old_segment[-30:]}")
                    
                    p.text = prefix + "(" + NEW_INTERACTIONS + "). " + suffix
                    print("Updated Interactions List.")
            except Exception as e:
                print(f"Error updating interactions: {e}")

def main():
    if not SOURCE_DOC.exists():
        print(f"Error: {SOURCE_DOC} missing.")
        return

    doc = Document(SOURCE_DOC)
    
    update_table(doc)
    update_text(doc)
    
    doc.save(OUT_DOC)
    print(f"--- Saved {OUT_DOC} ---")

if __name__ == "__main__":
    main()
